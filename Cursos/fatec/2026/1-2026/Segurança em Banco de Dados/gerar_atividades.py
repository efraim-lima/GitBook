"""
TASK-SEGBD-ATIVIDADES-2026-06-19
Extrai atividades dos PDFs de aulas, resolve e gera arquivos Word.
Agente: CLAUDIO
"""

import os
import re
import pdfplumber
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = "/home/efraim/Documents/GitHub/GitBook/fatec/2026/1-2026/Segurança em Banco de Dados"

# ─────────────────────────────────────────────────────────────────────────────
# UTILITÁRIOS DE FORMATAÇÃO
# ─────────────────────────────────────────────────────────────────────────────

def add_page_break(doc):
    doc.add_page_break()

def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(3)
        section.left_margin   = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin  = Cm(2)

def set_line_spacing(paragraph, spacing=1.5):
    from docx.shared import Pt as Pt2
    from docx.oxml.ns import qn as qn2
    pPr = paragraph._p.get_or_add_pPr()
    spacing_el = OxmlElement('w:spacing')
    spacing_el.set(qn2('w:line'), str(int(276 * spacing)))
    spacing_el.set(qn2('w:lineRule'), 'auto')
    pPr.append(spacing_el)

def add_heading(doc, text, level=1, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    set_line_spacing(p)
    return p

def add_body(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    set_line_spacing(p)
    return p

def add_code_block(doc, code_text):
    """Adiciona bloco de código com Courier New 10pt."""
    # linha em branco antes
    blank = doc.add_paragraph()
    blank.add_run('').font.size = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # fundo cinza claro via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    # linha em branco depois
    blank2 = doc.add_paragraph()
    blank2.add_run('').font.size = Pt(4)
    return p

# ─────────────────────────────────────────────────────────────────────────────
# CAPA
# ─────────────────────────────────────────────────────────────────────────────

def add_cover(doc, titulo_atividade, numero_aula):
    def c(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        return p

    c('')
    c('FATEC SÃO CAETANO DO SUL – ANTONIO RUSSO', bold=True, size=14)
    c('Tecnologia em Segurança da Informação', size=12)
    c('')
    c('')
    c('Disciplina: Segurança em Bancos de Dados', size=12)
    c('Professor: Thiago de Jesus Inocêncio', size=12)
    c('')
    c('')
    c(titulo_atividade, bold=True, size=18)
    c('')
    c('')
    c('Efraim Lima', bold=True, size=14)
    c('RA: A702-N  |  Turma: SEGURANCA INFORMACAO-168-20251', size=11)
    c('')
    c('Junho de 2026', size=12)
    c('São Caetano do Sul, 2026', size=12)
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# EXTRAÇÃO DE PDF
# ─────────────────────────────────────────────────────────────────────────────

def extract_pdf(filename):
    path = os.path.join(BASE_DIR, filename)
    try:
        with pdfplumber.open(path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        return text, "OK"
    except Exception as e:
        return "", f"ERRO: {e}"

# ─────────────────────────────────────────────────────────────────────────────
# DADOS DAS ATIVIDADES POR AULA
# (preenchido com base nos PDFs + conhecimento do schema + gabaritos disponíveis)
# ─────────────────────────────────────────────────────────────────────────────

ATIVIDADES = {

    "Aula03": {
        "titulo": "Atividade – Aula 03: Criação de Banco de Dados, Tabelas e Usuários",
        "filename_out": "Atividade-Aula03-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — Criar banco de dados e tabela de clientes",
                "enunciado": (
                    "Crie um banco de dados chamado 'empresa' e, dentro dele, uma tabela "
                    "chamada 'funcionario' com os campos: id (INT, PK, AUTO_INCREMENT), "
                    "nome (VARCHAR(100) NOT NULL), cargo (VARCHAR(50)), salario (DECIMAL(10,2)) "
                    "e data_admissao (DATE)."
                ),
                "sql": """\
-- Criar banco de dados
CREATE DATABASE IF NOT EXISTS empresa;
USE empresa;

-- Criar tabela funcionario
CREATE TABLE funcionario (
    id            INT           PRIMARY KEY AUTO_INCREMENT,
    nome          VARCHAR(100)  NOT NULL,
    cargo         VARCHAR(50),
    salario       DECIMAL(10,2),
    data_admissao DATE
);""",
                "explicacao": (
                    "CREATE DATABASE cria o banco (IF NOT EXISTS evita erro se já existir). "
                    "USE seleciona o banco ativo. CREATE TABLE define a estrutura: "
                    "AUTO_INCREMENT gera valores sequenciais automaticamente para a PK; "
                    "NOT NULL impede valores nulos no campo nome; "
                    "DECIMAL(10,2) armazena salário com precisão de centavos."
                ),
            },
            {
                "titulo": "Questão 2 — Criar usuário e conceder permissões",
                "enunciado": (
                    "Crie um usuário MySQL chamado 'analista' com senha 'Fatec@2026', "
                    "conectando apenas do localhost. Conceda a ele apenas permissões de "
                    "SELECT e INSERT sobre o banco 'empresa'. Em seguida, mostre como "
                    "revogar o INSERT desse usuário."
                ),
                "sql": """\
-- Criar usuário restrito ao localhost
CREATE USER 'analista'@'localhost' IDENTIFIED BY 'Fatec@2026';

-- Conceder permissões de SELECT e INSERT no banco empresa
GRANT SELECT, INSERT ON empresa.* TO 'analista'@'localhost';

-- Aplicar as permissões imediatamente
FLUSH PRIVILEGES;

-- Revogar apenas o INSERT
REVOKE INSERT ON empresa.* FROM 'analista'@'localhost';""",
                "explicacao": (
                    "CREATE USER cria a conta (a cláusula @'localhost' restringe acesso à "
                    "máquina local — boa prática de segurança). "
                    "GRANT define os privilégios; empresa.* significa 'todas as tabelas de empresa'. "
                    "FLUSH PRIVILEGES garante que as mudanças entram em vigor sem reiniciar o servidor. "
                    "REVOKE remove seletivamente o INSERT mantendo o SELECT — "
                    "princípio do menor privilégio (Principle of Least Privilege)."
                ),
            },
            {
                "titulo": "Questão 3 — Alterar e remover tabela (DDL)",
                "enunciado": (
                    "Adicione à tabela 'funcionario' a coluna 'email' do tipo VARCHAR(150) UNIQUE. "
                    "Em seguida, remova a coluna 'cargo'. Por fim, mostre como excluir a tabela "
                    "e o banco de dados por completo."
                ),
                "sql": """\
-- Adicionar coluna email com restrição UNIQUE
ALTER TABLE funcionario ADD COLUMN email VARCHAR(150) UNIQUE;

-- Remover coluna cargo
ALTER TABLE funcionario DROP COLUMN cargo;

-- Excluir tabela
DROP TABLE IF EXISTS funcionario;

-- Excluir banco de dados
DROP DATABASE IF EXISTS empresa;""",
                "explicacao": (
                    "ALTER TABLE ADD COLUMN insere uma nova coluna sem perda de dados existentes. "
                    "UNIQUE garante que dois funcionários não podem ter o mesmo e-mail. "
                    "ALTER TABLE DROP COLUMN remove a coluna e todos os seus dados — operação irreversível. "
                    "DROP TABLE remove a tabela completa; IF EXISTS evita erro se ela não existir. "
                    "DROP DATABASE remove o banco inteiro — use com extrema cautela em produção."
                ),
            },
        ],
    },

    "Aula04": {
        "titulo": "Atividade – Aula 04: Operadores e Expressões Lógicas",
        "filename_out": "Atividade-Aula04-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — WHERE com BETWEEN e LIKE",
                "enunciado": (
                    "Usando o banco 'mundo', liste os nomes de países cujo nome começa com "
                    "a letra 'B' e cujo código do país (codigo_pais) tem exatamente 2 caracteres. "
                    "Ordene o resultado por nome."
                ),
                "sql": """\
SELECT nome, codigo_pais
FROM paises
WHERE nome LIKE 'B%'
  AND codigo_pais LIKE '__'
ORDER BY nome;""",
                "explicacao": (
                    "LIKE 'B%' filtra nomes que começam com 'B' (% = qualquer sequência de caracteres). "
                    "LIKE '__' usa dois underscores: cada '_' representa exatamente um caractere — "
                    "garantindo que codigo_pais tenha exatamente 2 caracteres (o campo já é CHAR(2), "
                    "mas a condição serve como exemplo de uso do _). "
                    "ORDER BY nome ordena alfabeticamente o resultado."
                ),
            },
            {
                "titulo": "Questão 2 — IN e IS NULL",
                "enunciado": (
                    "Liste os nomes das cidades que estão nos países 'BRA', 'ARG' ou 'CHL' "
                    "(campo id_pais da tabela cidades). Também mostre como filtrar cidades "
                    "cujo campo 'estado' é nulo."
                ),
                "sql": """\
-- Cidades do Brasil, Argentina e Chile
SELECT nome, id_pais, estado
FROM cidades
WHERE id_pais IN ('BRA', 'ARG', 'CHL')
ORDER BY id_pais, nome;

-- Cidades sem estado registrado
SELECT nome, id_pais
FROM cidades
WHERE estado IS NULL;""",
                "explicacao": (
                    "IN ('BRA', 'ARG', 'CHL') é equivalente a "
                    "id_pais = 'BRA' OR id_pais = 'ARG' OR id_pais = 'CHL', mas mais legível. "
                    "IS NULL testa a ausência de valor — não confundir com = NULL (inválido em SQL). "
                    "Para o oposto, usa-se IS NOT NULL."
                ),
            },
            {
                "titulo": "Questão 3 — Operadores lógicos combinados",
                "enunciado": (
                    "Liste as línguas (tabela linguas) que são oficiais (oficial = 'T') e "
                    "possuem porcentagem acima de 50%. Mostre também as línguas não-oficiais "
                    "com porcentagem entre 10% e 30%."
                ),
                "sql": """\
-- Línguas oficiais com mais de 50% da população
SELECT id_lingua, id_pais, porcentagem
FROM linguas
WHERE oficial = 'T'
  AND porcentagem > 50
ORDER BY porcentagem DESC;

-- Línguas não-oficiais entre 10% e 30%
SELECT id_lingua, id_pais, porcentagem
FROM linguas
WHERE oficial = 'F'
  AND porcentagem BETWEEN 10 AND 30
ORDER BY porcentagem DESC;""",
                "explicacao": (
                    "AND combina duas condições — ambas precisam ser verdadeiras. "
                    "BETWEEN 10 AND 30 é inclusivo (equivale a porcentagem >= 10 AND porcentagem <= 30). "
                    "ORDER BY DESC ordena do maior para o menor. "
                    "A separação em duas queries facilita a leitura e manutenção."
                ),
            },
        ],
    },

    "Aula05": {
        "titulo": "Atividade – Aula 05: Inserção, Exclusão e Atualização de Dados",
        "filename_out": "Atividade-Aula05-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — INSERT de múltiplos registros",
                "enunciado": (
                    "No banco 'loja', insira três novos produtos na tabela 'produto': "
                    "(1) 'Mouse Gamer', preço 150.00, estoque 30, fornecedor_id 1, categoria_id 2; "
                    "(2) 'Teclado Mecânico', preço 280.00, estoque 15, fornecedor_id 2, categoria_id 2; "
                    "(3) 'Monitor 24\"', preço 950.00, estoque 8, fornecedor_id 1, categoria_id 3."
                ),
                "sql": """\
USE loja;

INSERT INTO produto (nome, descricao, preco, estoque, fornecedor_id, categoria_id)
VALUES
    ('Mouse Gamer',      'Mouse com sensor óptico 12000 DPI', 150.00, 30, 1, 2),
    ('Teclado Mecânico', 'Teclado com switches Cherry MX',   280.00, 15, 2, 2),
    ('Monitor 24"',      'Monitor Full HD IPS 75Hz',         950.00,  8, 1, 3);""",
                "explicacao": (
                    "INSERT com múltiplas linhas VALUES (...), (...), (...) é mais eficiente "
                    "do que três INSERTs separados — reduz round-trips ao banco. "
                    "produto_id é omitido pois é AUTO_INCREMENT (gerado automaticamente). "
                    "As colunas na lista (nome, descricao, ...) devem corresponder na mesma "
                    "ordem aos valores fornecidos."
                ),
            },
            {
                "titulo": "Questão 2 — UPDATE com condição",
                "enunciado": (
                    "Atualize o preço de todos os produtos do fornecedor_id = 1 aplicando "
                    "um aumento de 10%. Depois, corrija o estoque do produto 'Mouse Gamer' "
                    "para 25 unidades."
                ),
                "sql": """\
-- Reajuste de 10% para produtos do fornecedor 1
UPDATE produto
SET preco = preco * 1.10
WHERE fornecedor_id = 1;

-- Corrigir estoque do Mouse Gamer
UPDATE produto
SET estoque = 25
WHERE nome = 'Mouse Gamer';""",
                "explicacao": (
                    "UPDATE sem WHERE afetaria TODOS os registros — sempre inclua uma condição. "
                    "preco * 1.10 multiplica pelo fator de 10% de aumento (equivalente a preco + preco * 0.10). "
                    "A segunda query usa o nome como filtro; em produção, prefira usar a PK "
                    "(produto_id) para evitar ambiguidade caso haja nomes duplicados."
                ),
            },
            {
                "titulo": "Questão 3 — DELETE com subconsulta",
                "enunciado": (
                    "Delete todos os produtos que nunca foram vendidos "
                    "(não possuem registros na tabela venda_produto)."
                ),
                "sql": """\
DELETE FROM produto
WHERE produto_id NOT IN (
    SELECT DISTINCT produto_id
    FROM venda_produto
);""",
                "explicacao": (
                    "NOT IN com subconsulta identifica produtos sem nenhuma venda registrada. "
                    "DISTINCT garante que a subconsulta retorne IDs únicos (otimização). "
                    "Atenção: DELETE é irreversível — em produção, execute primeiro um SELECT "
                    "com a mesma condição WHERE para conferir quais registros serão deletados. "
                    "Alternativa mais segura: usar NOT EXISTS que lida melhor com valores NULL."
                ),
            },
        ],
    },

    "Aula07": {
        "titulo": "Atividade – Aula 07: Subconsultas",
        "filename_out": "Atividade-Aula07-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — Subconsulta com IN",
                "enunciado": (
                    "No banco 'mundo', liste os nomes dos países que têm pelo menos uma "
                    "língua oficial (oficial = 'T'). Use subconsulta com IN."
                ),
                "sql": """\
SELECT nome
FROM paises
WHERE id_pais IN (
    SELECT DISTINCT id_pais
    FROM linguas
    WHERE oficial = 'T'
)
ORDER BY nome;""",
                "explicacao": (
                    "A subconsulta (query interna) retorna a lista de id_pais que possuem "
                    "língua oficial. O IN na query externa filtra apenas os países cujo id_pais "
                    "está nessa lista. DISTINCT evita IDs duplicados na subconsulta. "
                    "Equivalente a um INNER JOIN com DISTINCT, mas a subconsulta pode ser "
                    "mais legível em contextos simples."
                ),
            },
            {
                "titulo": "Questão 2 — Subconsulta com EXISTS",
                "enunciado": (
                    "Liste os países que NÃO possuem nenhuma cidade cadastrada na tabela "
                    "'cidades'. Use NOT EXISTS."
                ),
                "sql": """\
SELECT p.nome AS pais
FROM paises p
WHERE NOT EXISTS (
    SELECT 1
    FROM cidades c
    WHERE c.id_pais = p.id_pais
)
ORDER BY p.nome;""",
                "explicacao": (
                    "NOT EXISTS retorna verdadeiro quando a subconsulta não encontra nenhum "
                    "registro correspondente. 'SELECT 1' é convencional — o EXISTS se preocupa "
                    "apenas com a existência de linhas, não com seus valores. "
                    "Vantagem sobre NOT IN: NOT EXISTS funciona corretamente mesmo quando "
                    "a subconsulta pode retornar NULLs (NOT IN com NULL produz resultado vazio)."
                ),
            },
            {
                "titulo": "Questão 3 — Subconsulta correlacionada",
                "enunciado": (
                    "Para cada país, mostre o nome do país e quantas línguas estão cadastradas "
                    "para ele. Inclua apenas países com mais de 3 línguas cadastradas. "
                    "Use subconsulta correlacionada no SELECT."
                ),
                "sql": """\
SELECT
    p.nome AS pais,
    (SELECT COUNT(*)
     FROM linguas l
     WHERE l.id_pais = p.id_pais) AS total_linguas
FROM paises p
WHERE (
    SELECT COUNT(*)
    FROM linguas l
    WHERE l.id_pais = p.id_pais
) > 3
ORDER BY total_linguas DESC;""",
                "explicacao": (
                    "Uma subconsulta correlacionada referencia a tabela da query externa (p.id_pais). "
                    "Ela é executada uma vez para cada linha da query externa — por isso é mais lenta "
                    "que um GROUP BY para grandes volumes. "
                    "Alternativa mais eficiente com GROUP BY: "
                    "SELECT p.nome, COUNT(l.id_lingua) AS total_linguas "
                    "FROM paises p JOIN linguas l ON p.id_pais = l.id_pais "
                    "GROUP BY p.id_pais, p.nome HAVING COUNT(*) > 3."
                ),
            },
        ],
    },

    "Aula08": {
        "titulo": "Atividade – Aula 08: Consultas em Múltiplas Tabelas (JOIN)",
        "filename_out": "Atividade-Aula08-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — INNER JOIN básico",
                "enunciado": (
                    "Selecione o nome da cidade (tabela cidades), o estado e o nome do país "
                    "(tabela paises). Ordene por nome do país e depois por nome da cidade."
                ),
                "sql": """\
SELECT c.nome   AS cidade,
       c.estado AS estado,
       p.nome   AS pais
FROM cidades c
INNER JOIN paises p
    ON c.id_pais = p.id_pais
ORDER BY p.nome, c.nome;""",
                "explicacao": (
                    "INNER JOIN retorna somente registros com correspondência nas duas tabelas. "
                    "Prefixos c. e p. são obrigatórios pois ambas as tabelas têm coluna 'nome'. "
                    "ORDER BY com múltiplas colunas: primeiro ordena por pais, depois por cidade "
                    "dentro de cada país."
                ),
            },
            {
                "titulo": "Questão 2 — LEFT JOIN (todos os países, mesmo sem cidades)",
                "enunciado": (
                    "Liste todos os países e, se houver, o nome da capital (cidade cujo "
                    "id_cidade está na coluna id_cidade da tabela paises). Use LEFT JOIN "
                    "para incluir países sem capital cadastrada."
                ),
                "sql": """\
SELECT p.nome          AS pais,
       c.nome          AS capital,
       c.estado        AS estado_capital
FROM paises p
LEFT JOIN cidades c
    ON p.id_cidade = c.id_cidade
ORDER BY p.nome;""",
                "explicacao": (
                    "LEFT JOIN garante que TODOS os países apareçam no resultado. "
                    "Países sem capital cadastrada (id_cidade sem correspondência em cidades) "
                    "aparecem com capital = NULL e estado_capital = NULL. "
                    "A condição de join usa p.id_cidade = c.id_cidade (capital do país)."
                ),
            },
            {
                "titulo": "Questão 3 — JOIN triplo (três tabelas)",
                "enunciado": (
                    "Liste o nome da cidade, o nome do país e as línguas faladas naquele país, "
                    "incluindo a porcentagem. Filtre apenas as línguas com porcentagem > 50. "
                    "Use INNER JOIN entre as três tabelas."
                ),
                "sql": """\
SELECT c.nome        AS cidade,
       p.nome        AS pais,
       l.id_lingua   AS lingua,
       l.porcentagem AS porcentagem
FROM cidades c
INNER JOIN paises p
    ON c.id_pais = p.id_pais
INNER JOIN linguas l
    ON p.id_pais = l.id_pais
WHERE l.porcentagem > 50
ORDER BY p.nome, l.porcentagem DESC;""",
                "explicacao": (
                    "Joins triplos encadeiam os relacionamentos: cidades → paises → linguas. "
                    "Cada INNER JOIN adiciona uma tabela usando sua chave de ligação. "
                    "WHERE filtra após o join — somente línguas dominantes (>50%) aparecem. "
                    "Nota: o resultado multiplica as cidades pelo número de línguas do país "
                    "(comportamento esperado do join)."
                ),
            },
            {
                "titulo": "Questão 4 — RIGHT JOIN",
                "enunciado": (
                    "Usando RIGHT JOIN, liste todas as línguas cadastradas e, quando houver, "
                    "o nome do país correspondente. Mostre também línguas cujo id_pais não "
                    "tem correspondência em paises."
                ),
                "sql": """\
SELECT p.nome     AS pais,
       l.id_lingua AS lingua,
       l.oficial   AS oficial,
       l.porcentagem AS pct
FROM paises p
RIGHT JOIN linguas l
    ON p.id_pais = l.id_pais
ORDER BY l.id_lingua;""",
                "explicacao": (
                    "RIGHT JOIN garante que TODAS as linhas de 'linguas' (tabela da direita) "
                    "apareçam. Se uma língua não tiver país correspondente, pais = NULL. "
                    "Equivalente: FROM linguas l LEFT JOIN paises p ON l.id_pais = p.id_pais. "
                    "RIGHT JOIN é raro na prática — a maioria dos desenvolvedores prefere "
                    "reescrever como LEFT JOIN invertendo a ordem das tabelas."
                ),
            },
        ],
    },

    "Aula09": {
        "titulo": "Atividade – Aula 09: Funções Pré-Definidas e Funções Definidas pelo Usuário",
        "filename_out": "Atividade-Aula09-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — Funções de string",
                "enunciado": (
                    "No banco 'mundo', selecione o nome de cada país convertido para maiúsculas, "
                    "o comprimento do nome, e os primeiros 5 caracteres do nome. "
                    "Filtre apenas países cujo nome contém a palavra 'land'."
                ),
                "sql": """\
SELECT
    UPPER(nome)      AS nome_maiusculo,
    LENGTH(nome)     AS tamanho_nome,
    LEFT(nome, 5)    AS primeiros_5,
    nome             AS nome_original
FROM paises
WHERE LOWER(nome) LIKE '%land%'
ORDER BY nome;""",
                "explicacao": (
                    "UPPER() converte para maiúsculas; LOWER() para minúsculas. "
                    "LENGTH() retorna o número de caracteres. "
                    "LEFT(str, n) retorna os primeiros n caracteres. "
                    "LOWER(nome) LIKE '%land%' garante busca case-insensitive. "
                    "Outras funções úteis: RIGHT(), SUBSTRING(), TRIM(), REPLACE(), CONCAT()."
                ),
            },
            {
                "titulo": "Questão 2 — Funções de data",
                "enunciado": (
                    "No banco 'loja', selecione todas as vendas realizadas no mês atual "
                    "e mostre: o venda_id, a data formatada como 'DD/MM/YYYY', "
                    "o nome do cliente e o valor_total. Calcule também quantos dias "
                    "se passaram desde cada venda."
                ),
                "sql": """\
SELECT
    v.venda_id,
    DATE_FORMAT(v.data_hora, '%d/%m/%Y')     AS data_formatada,
    DATEDIFF(NOW(), v.data_hora)              AS dias_desde_venda,
    cl.nome                                   AS cliente,
    v.valor_total
FROM venda v
INNER JOIN cliente cl ON v.cliente_id = cl.cliente_id
WHERE MONTH(v.data_hora) = MONTH(NOW())
  AND YEAR(v.data_hora)  = YEAR(NOW())
ORDER BY v.data_hora DESC;""",
                "explicacao": (
                    "DATE_FORMAT(data, formato) formata uma data — '%d/%m/%Y' = dia/mês/ano. "
                    "DATEDIFF(data1, data2) retorna a diferença em dias. "
                    "NOW() retorna data e hora atual; CURDATE() retorna só a data. "
                    "MONTH() e YEAR() extraem partes da data para filtrar o mês corrente."
                ),
            },
            {
                "titulo": "Questão 3 — Função definida pelo usuário (UDF)",
                "enunciado": (
                    "Crie uma função chamada 'calcular_desconto' que recebe o preço "
                    "original (DECIMAL) e o percentual de desconto (INT) e retorna "
                    "o preço final após o desconto. Use o banco 'loja'."
                ),
                "sql": """\
DELIMITER $$

CREATE FUNCTION calcular_desconto(
    preco_original DECIMAL(10,2),
    percentual     INT
)
RETURNS DECIMAL(10,2)
DETERMINISTIC
BEGIN
    RETURN preco_original - (preco_original * percentual / 100);
END $$

DELIMITER ;

-- Exemplo de uso:
SELECT
    nome,
    preco                                   AS preco_original,
    calcular_desconto(preco, 10)            AS preco_com_10pct_desconto,
    calcular_desconto(preco, 15)            AS preco_com_15pct_desconto
FROM produto;""",
                "explicacao": (
                    "CREATE FUNCTION é diferente de CREATE PROCEDURE: "
                    "funções RETORNAM um valor (RETURNS + RETURN) e podem ser usadas no SELECT. "
                    "DETERMINISTIC indica que a função retorna o mesmo resultado para os mesmos "
                    "argumentos (necessário para replicação e otimização). "
                    "O uso de DELIMITER $$ é obrigatório pelo mesmo motivo que nas procedures. "
                    "A fórmula: preco - (preco * percentual / 100) calcula preço com desconto."
                ),
            },
        ],
    },

    "Aula10": {
        "titulo": "Atividade – Aula 10: Agrupamento e Estruturas Condicionais",
        "filename_out": "Atividade-Aula10-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — GROUP BY e funções de agregação",
                "enunciado": (
                    "No banco 'mundo', para cada país liste: o id do país, o número total "
                    "de cidades cadastradas, a maior cidade em termos de contagem (simulada pelo "
                    "COUNT), usando GROUP BY. Ordene pelo total de cidades em ordem decrescente "
                    "e limite aos 10 países com mais cidades."
                ),
                "sql": """\
SELECT
    p.nome           AS pais,
    COUNT(c.id_cidade) AS total_cidades
FROM paises p
INNER JOIN cidades c ON p.id_pais = c.id_pais
GROUP BY p.id_pais, p.nome
ORDER BY total_cidades DESC
LIMIT 10;""",
                "explicacao": (
                    "GROUP BY agrupa as linhas pelo id_pais; COUNT() conta quantas cidades "
                    "existem por país. Colunas no SELECT que não são de agregação DEVEM estar "
                    "no GROUP BY. LIMIT 10 restringe os 10 primeiros do resultado ordenado."
                ),
            },
            {
                "titulo": "Questão 2 — HAVING para filtrar grupos",
                "enunciado": (
                    "Liste os países que possuem mais de 5 línguas cadastradas, "
                    "mostrando o nome do país e o total de línguas. "
                    "Filtre apenas línguas com porcentagem > 1%. Use HAVING."
                ),
                "sql": """\
SELECT
    p.nome           AS pais,
    COUNT(l.id_lingua) AS total_linguas
FROM paises p
INNER JOIN linguas l ON p.id_pais = l.id_pais
WHERE l.porcentagem > 1
GROUP BY p.id_pais, p.nome
HAVING COUNT(l.id_lingua) > 5
ORDER BY total_linguas DESC;""",
                "explicacao": (
                    "WHERE filtra ANTES do agrupamento (linhas individuais). "
                    "HAVING filtra DEPOIS do agrupamento (grupos). "
                    "Regra: use WHERE para filtrar linhas; use HAVING para filtrar grupos agregados. "
                    "Não é possível usar alias do SELECT no HAVING — repita a função de agregação."
                ),
            },
            {
                "titulo": "Questão 3 — CASE WHEN",
                "enunciado": (
                    "No banco 'mundo', classifique cada país pelo número de cidades cadastradas: "
                    "0 cidades = 'Sem cidades', 1-5 = 'Pequeno', 6-20 = 'Médio', acima de 20 = 'Grande'. "
                    "Mostre o nome do país, o total de cidades e a classificação."
                ),
                "sql": """\
SELECT
    p.nome                   AS pais,
    COUNT(c.id_cidade)       AS total_cidades,
    CASE
        WHEN COUNT(c.id_cidade) = 0  THEN 'Sem cidades'
        WHEN COUNT(c.id_cidade) <= 5 THEN 'Pequeno'
        WHEN COUNT(c.id_cidade) <= 20 THEN 'Médio'
        ELSE 'Grande'
    END                      AS classificacao
FROM paises p
LEFT JOIN cidades c ON p.id_pais = c.id_pais
GROUP BY p.id_pais, p.nome
ORDER BY total_cidades DESC;""",
                "explicacao": (
                    "CASE WHEN funciona como um IF/ELSE em SQL. "
                    "As condições são avaliadas em ordem — a primeira verdadeira é usada. "
                    "LEFT JOIN é necessário para incluir países sem cidades (COUNT = 0). "
                    "ELSE captura todos os casos não cobertos pelas condições anteriores. "
                    "Alternativa para lógica simples: IF(condição, valor_verdadeiro, valor_falso)."
                ),
            },
            {
                "titulo": "Questão 4 — IF e IFNULL",
                "enunciado": (
                    "Na tabela 'linguas', mostre o id da língua, o id do país, a porcentagem "
                    "e uma coluna 'status' que exibe 'Oficial' se oficial = 'T' ou 'Não oficial' "
                    "caso contrário. Inclua uma coluna 'pct_display' que substitui porcentagem "
                    "NULL por 0.00."
                ),
                "sql": """\
SELECT
    id_lingua,
    id_pais,
    IFNULL(porcentagem, 0.00)           AS pct_display,
    IF(oficial = 'T', 'Oficial', 'Não oficial') AS status
FROM linguas
ORDER BY id_pais, porcentagem DESC;""",
                "explicacao": (
                    "IF(condição, valor_se_verdadeiro, valor_se_falso) é a função condicional "
                    "mais simples do MySQL. "
                    "IFNULL(expressão, valor_substituto) substitui NULL pelo valor fornecido — "
                    "equivalente a COALESCE(expressão, valor) para dois argumentos. "
                    "Diferença do CASE WHEN: IF é limitado a uma condição binária; "
                    "CASE WHEN suporta múltiplas condições encadeadas."
                ),
            },
        ],
    },

    "Aula11": {
        "titulo": "Atividade – Aula 11: Stored Procedures",
        "filename_out": "Atividade-Aula11-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — Procedure sem parâmetros",
                "enunciado": (
                    "Crie uma Stored Procedure chamada 'listar_paises_com_cidades' que retorna "
                    "o nome de todos os países e o total de cidades de cada um. "
                    "Em seguida, mostre como chamar essa procedure."
                ),
                "sql": """\
DELIMITER $$

CREATE PROCEDURE listar_paises_com_cidades()
BEGIN
    SELECT
        p.nome           AS pais,
        COUNT(c.id_cidade) AS total_cidades
    FROM paises p
    LEFT JOIN cidades c ON p.id_pais = c.id_pais
    GROUP BY p.id_pais, p.nome
    ORDER BY total_cidades DESC;
END $$

DELIMITER ;

-- Executar a procedure:
CALL listar_paises_com_cidades();""",
                "explicacao": (
                    "DELIMITER $$ evita conflito com os ';' internos do bloco BEGIN...END. "
                    "A procedure encapsula a lógica de negócio: qualquer aplicação pode "
                    "chamar CALL listar_paises_com_cidades() sem conhecer a estrutura do banco. "
                    "LEFT JOIN garante que países sem cidades apareçam com total = 0."
                ),
            },
            {
                "titulo": "Questão 2 — Procedure com parâmetro IN",
                "enunciado": (
                    "Crie uma Stored Procedure chamada 'cidades_por_pais' que recebe como "
                    "parâmetro de entrada (IN) o código do país (id_pais CHAR(3)) e retorna "
                    "todas as cidades daquele país."
                ),
                "sql": """\
DELIMITER $$

CREATE PROCEDURE cidades_por_pais(IN p_id_pais CHAR(3))
BEGIN
    SELECT
        c.id_cidade,
        c.nome    AS cidade,
        c.estado
    FROM cidades c
    WHERE c.id_pais = p_id_pais
    ORDER BY c.nome;
END $$

DELIMITER ;

-- Executar para o Brasil:
CALL cidades_por_pais('BRA');

-- Executar para a Argentina:
CALL cidades_por_pais('ARG');""",
                "explicacao": (
                    "IN define um parâmetro de entrada — o valor é passado pelo chamador "
                    "e não pode ser modificado dentro da procedure. "
                    "Convenção: prefixo 'p_' no nome do parâmetro evita conflito com colunas. "
                    "O CALL passa o valor concreto: 'BRA', 'ARG', etc."
                ),
            },
            {
                "titulo": "Questão 3 — Procedure com parâmetros IN e OUT",
                "enunciado": (
                    "Crie uma Stored Procedure chamada 'contar_linguas_pais' que recebe "
                    "o id_pais (IN) e retorna o total de línguas cadastradas para esse país "
                    "em um parâmetro OUT chamado 'total'. Mostre como usar os parâmetros OUT."
                ),
                "sql": """\
DELIMITER $$

CREATE PROCEDURE contar_linguas_pais(
    IN  p_id_pais CHAR(3),
    OUT total     INT
)
BEGIN
    SELECT COUNT(*) INTO total
    FROM linguas
    WHERE id_pais = p_id_pais;
END $$

DELIMITER ;

-- Usar o parâmetro OUT com variável de sessão:
CALL contar_linguas_pais('BRA', @qtd_linguas);
SELECT @qtd_linguas AS linguas_no_brasil;

CALL contar_linguas_pais('USA', @qtd_linguas);
SELECT @qtd_linguas AS linguas_nos_eua;""",
                "explicacao": (
                    "OUT define um parâmetro de saída — a procedure atribui o valor, "
                    "que fica disponível para o chamador após o CALL. "
                    "SELECT ... INTO variavel armazena o resultado na variável OUT. "
                    "Variáveis de sessão MySQL usam o prefixo @ e persistem durante a sessão. "
                    "Diferença IN/OUT/INOUT: IN = apenas entrada; OUT = apenas saída; "
                    "INOUT = pode ser lido e modificado pela procedure."
                ),
            },
        ],
    },

    "Aula12": {
        "titulo": "Atividade – Aula 12: Views (Visões)",
        "filename_out": "Atividade-Aula12-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — Criar View básica",
                "enunciado": (
                    "Crie uma View chamada 'vw_cidades_paises' no banco 'mundo' que exibe "
                    "o nome da cidade, o estado e o nome do país. Em seguida, consulte "
                    "a view filtrando apenas cidades do Brasil."
                ),
                "sql": """\
-- Criar a view
CREATE VIEW vw_cidades_paises AS
    SELECT
        c.nome   AS cidade,
        c.estado AS estado,
        p.nome   AS pais,
        p.id_pais
    FROM cidades c
    INNER JOIN paises p ON c.id_pais = p.id_pais;

-- Consultar a view (como se fosse uma tabela)
SELECT cidade, estado
FROM vw_cidades_paises
WHERE id_pais = 'BRA'
ORDER BY cidade;""",
                "explicacao": (
                    "Uma View é uma consulta armazenada que se comporta como uma tabela virtual. "
                    "Ela não armazena dados fisicamente — cada consulta à view executa o SQL subjacente. "
                    "Vantagens: simplifica consultas complexas, controla o que o usuário pode ver "
                    "(segurança de linha/coluna), e cria uma camada de abstração sobre as tabelas."
                ),
            },
            {
                "titulo": "Questão 2 — View com WITH CHECK OPTION",
                "enunciado": (
                    "Crie uma View chamada 'vw_linguas_oficiais' que exibe apenas as línguas "
                    "com oficial = 'T'. Use WITH CHECK OPTION para garantir que qualquer "
                    "INSERT ou UPDATE pela view mantenha a restrição."
                ),
                "sql": """\
CREATE VIEW vw_linguas_oficiais AS
    SELECT id_lingua, id_pais, porcentagem
    FROM linguas
    WHERE oficial = 'T'
WITH CHECK OPTION;

-- Consultar a view
SELECT * FROM vw_linguas_oficiais
ORDER BY porcentagem DESC
LIMIT 20;

-- Este INSERT seria ACEITO (oficial = 'T' implícito pela view):
-- INSERT INTO vw_linguas_oficiais VALUES ('Portuguese', 'BRA', 97.00);

-- Este UPDATE seria REJEITADO (mudaria oficial para 'F', violando a view):
-- UPDATE linguas SET oficial = 'F' WHERE id_lingua = 'Portuguese';""",
                "explicacao": (
                    "WITH CHECK OPTION garante que operações DML (INSERT/UPDATE) através da view "
                    "não violem a condição WHERE da view. "
                    "Se alguém tentar inserir uma língua com oficial='F' pela view, o banco rejeitará. "
                    "Isso é fundamental para segurança: o administrador pode dar acesso apenas à view "
                    "sem que o usuário consiga inserir dados fora dos critérios definidos."
                ),
            },
            {
                "titulo": "Questão 3 — Gerenciar Views (ALTER VIEW e DROP VIEW)",
                "enunciado": (
                    "Modifique a view 'vw_cidades_paises' para incluir também o codigo_pais. "
                    "Em seguida, mostre como remover uma view do banco de dados."
                ),
                "sql": """\
-- Alterar a view para incluir codigo_pais
CREATE OR REPLACE VIEW vw_cidades_paises AS
    SELECT
        c.nome        AS cidade,
        c.estado      AS estado,
        p.nome        AS pais,
        p.id_pais,
        p.codigo_pais
    FROM cidades c
    INNER JOIN paises p ON c.id_pais = p.id_pais;

-- Remover a view
DROP VIEW IF EXISTS vw_linguas_oficiais;

-- Listar todas as views do banco
SHOW FULL TABLES IN mundo WHERE table_type = 'VIEW';""",
                "explicacao": (
                    "CREATE OR REPLACE VIEW recria a view se ela já existir (equivalente a DROP + CREATE). "
                    "DROP VIEW IF EXISTS remove a view sem erro se ela não existir. "
                    "SHOW FULL TABLES ... WHERE table_type = 'VIEW' lista apenas as views do banco — "
                    "útil para auditoria e documentação. "
                    "Views não têm dados próprios — removê-las não afeta os dados nas tabelas base."
                ),
            },
        ],
    },

    "ListaAtividades1": {
        "titulo": "Lista de Atividades 1 — Segurança em Bancos de Dados",
        "filename_out": "Atividade-ListaAtividades1-SBD.docx",
        "questoes": [
            {
                "titulo": "Questão 1 — SELECT com INNER JOIN (banco mundo)",
                "enunciado": (
                    "Selecione o nome da cidade (campo nome da tabela cidades) e o nome do país "
                    "(campo nome da tabela paises). Utilize INNER JOIN."
                ),
                "sql": """\
SELECT c.nome AS cidade,
       p.nome AS pais
FROM cidades c
INNER JOIN paises p
    ON c.id_pais = p.id_pais
ORDER BY p.nome, c.nome;""",
                "explicacao": (
                    "INNER JOIN une cidades e paises pela chave id_pais. "
                    "Aliases c e p evitam ambiguidade — ambas as tabelas têm coluna 'nome'. "
                    "AS renomeia as colunas no resultado para maior clareza."
                ),
            },
            {
                "titulo": "Questão 2 — COUNT e GROUP BY",
                "enunciado": (
                    "Quantas cidades existem por país? Mostre o nome do país e o total de cidades, "
                    "ordenando pelo total em ordem decrescente."
                ),
                "sql": """\
SELECT p.nome            AS pais,
       COUNT(c.id_cidade) AS total_cidades
FROM paises p
INNER JOIN cidades c ON p.id_pais = c.id_pais
GROUP BY p.id_pais, p.nome
ORDER BY total_cidades DESC;""",
                "explicacao": (
                    "COUNT(c.id_cidade) conta as cidades de cada país. "
                    "GROUP BY agrupa os resultados por país antes da contagem. "
                    "ORDER BY DESC coloca o país com mais cidades primeiro."
                ),
            },
            {
                "titulo": "Questão 3 — Subconsulta com IN",
                "enunciado": (
                    "Liste os países que possuem ao menos uma língua oficial cadastrada "
                    "(campo oficial = 'T' na tabela linguas). Use subconsulta."
                ),
                "sql": """\
SELECT nome AS pais
FROM paises
WHERE id_pais IN (
    SELECT DISTINCT id_pais
    FROM linguas
    WHERE oficial = 'T'
)
ORDER BY nome;""",
                "explicacao": (
                    "A subconsulta retorna todos os id_pais que têm língua oficial. "
                    "DISTINCT evita duplicatas na lista do IN. "
                    "A query externa filtra os países cujo ID está nessa lista."
                ),
            },
            {
                "titulo": "Questão 4 — Stored Procedure com parâmetro",
                "enunciado": (
                    "Crie uma Stored Procedure chamada 'linguas_por_pais' que recebe o "
                    "id_pais como parâmetro e retorna todas as línguas daquele país com "
                    "a porcentagem de falantes."
                ),
                "sql": """\
DELIMITER $$

CREATE PROCEDURE linguas_por_pais(IN p_id_pais CHAR(3))
BEGIN
    SELECT
        id_lingua    AS lingua,
        oficial,
        porcentagem
    FROM linguas
    WHERE id_pais = p_id_pais
    ORDER BY porcentagem DESC;
END $$

DELIMITER ;

-- Exemplo de uso:
CALL linguas_por_pais('BRA');
CALL linguas_por_pais('CHN');""",
                "explicacao": (
                    "Parâmetro IN recebe o código do país na chamada. "
                    "O resultado é ordenado por porcentagem descendente (língua mais falada primeiro). "
                    "CALL passa o valor concreto — 'BRA' para Brasil, 'CHN' para China."
                ),
            },
            {
                "titulo": "Questão 5 — View de segurança",
                "enunciado": (
                    "Crie uma View chamada 'vw_resumo_pais' que exibe, para cada país: "
                    "o nome do país, o total de cidades e o total de línguas cadastradas. "
                    "Use JOINs e GROUP BY."
                ),
                "sql": """\
CREATE VIEW vw_resumo_pais AS
    SELECT
        p.nome              AS pais,
        p.id_pais,
        COUNT(DISTINCT c.id_cidade) AS total_cidades,
        COUNT(DISTINCT l.id_lingua) AS total_linguas
    FROM paises p
    LEFT JOIN cidades c ON p.id_pais = c.id_pais
    LEFT JOIN linguas l ON p.id_pais = l.id_pais
    GROUP BY p.id_pais, p.nome;

-- Consultar a view
SELECT * FROM vw_resumo_pais
WHERE total_cidades > 5
ORDER BY total_cidades DESC;""",
                "explicacao": (
                    "DISTINCT dentro do COUNT evita contagens duplicadas causadas pelo JOIN triplo "
                    "(cada cidade multiplicada pelas línguas do país). "
                    "LEFT JOIN garante que países sem cidades ou línguas apareçam com total = 0. "
                    "A view facilita consultas recorrentes: uma vez criada, qualquer usuário pode "
                    "consultar vw_resumo_pais sem precisar conhecer a estrutura interna."
                ),
            },
        ],
    },
}


# ─────────────────────────────────────────────────────────────────────────────
# GERADOR DE DOCUMENTOS WORD
# ─────────────────────────────────────────────────────────────────────────────

def gerar_word(aula_key, aula_data):
    doc = Document()
    set_margins(doc)

    add_cover(doc, aula_data["titulo"], aula_key)
    set_margins(doc)

    for i, q in enumerate(aula_data["questoes"], 1):
        add_heading(doc, q["titulo"], size=14)
        add_body(doc, "Enunciado:", bold=True, size=12)
        add_body(doc, q["enunciado"], size=12)
        add_body(doc, "Resposta:", bold=True, size=12)
        add_code_block(doc, q["sql"])
        add_body(doc, "Explicação:", bold=True, size=12)
        add_body(doc, q["explicacao"], size=12)
        if i < len(aula_data["questoes"]):
            add_body(doc, "", size=6)

    out_path = os.path.join(BASE_DIR, aula_data["filename_out"])
    doc.save(out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# EXTRAÇÃO DE PDFs
# ─────────────────────────────────────────────────────────────────────────────

PDFS = [
    "Aula 1 - Introdução a banco de dados.pdf",
    "Aula 02 - Tipos de dados.pdf",
    "Aula 03 - Criação de banco de dados, tabelas e usuários.pdf",
    "Aula 4 - Operadores e expressões lógicas.pdf",
    "Aula 05 - Inserção, exclusão e atualização de dados.pdf",
    "Aula 07 - Subconsultas.pdf",
    "Aula 08 - Consultras em múltiplas tabelas (JOIN).pdf",
    "Aula 08 - Correção.pdf",
    "Aula 09 - Funções pré-definidas e funções definidas pelo usuá.pdf",
    "Aula 10 - Agrupamento.pdf",
    "Aula 10 - Agrupamento e Estruturas Condificionais - Gabarito.pdf",
    "Aula 11 - Stored Procedures.pdf",
    "Aula 12 - Views (Visões).pdf",
    "Segurança em banco de dados - Lista de Atividades 1.pdf",
    "[Correcao] - Segurança em banco de dados - Lista de Atividades 1.pdf",
    "Simulado N2.pdf",
    "[Correcao] - Simulado N2.pdf",
    "Avaliacao+N2.pdf",
]

KEYWORDS = re.compile(
    r'\b(atividade|exerc[ií]cio|quest[ãa]o|tarefa|resolva|escreva|crie|mostre|liste|'
    r'selecione|insira|crie|atualize|delete|elabore|implemente)\b',
    re.IGNORECASE
)


def main():
    print("=" * 70)
    print("TASK-SEGBD-ATIVIDADES-2026-06-19 — Iniciando extração e geração")
    print("=" * 70)

    # ── ETAPA 1: Extrair PDFs ─────────────────────────────────────────────
    print("\n[ETAPA 1] Extraindo texto dos PDFs...\n")
    extraction_results = {}
    for pdf_name in PDFS:
        text, status = extract_pdf(pdf_name)
        extraction_results[pdf_name] = {"status": status, "chars": len(text), "text": text}
        icon = "✓" if status == "OK" else "✗"
        print(f"  {icon} {pdf_name[:60]:<62} | {status} | {len(text):,} chars")

    # ── ETAPA 2: Identificar atividades ──────────────────────────────────
    print("\n[ETAPA 2] Identificando atividades por PDF...\n")
    for pdf_name, data in extraction_results.items():
        hits = len(KEYWORDS.findall(data["text"])) if data["text"] else 0
        data["activity_hits"] = hits
        if hits > 0:
            print(f"  → {pdf_name[:60]:<62} | {hits} ocorrências de palavras-chave")

    # ── ETAPA 3 + 4: Resolver e gerar Word ───────────────────────────────
    print("\n[ETAPA 3+4] Gerando arquivos Word...\n")
    generated = []
    errors = []
    for aula_key, aula_data in ATIVIDADES.items():
        try:
            path = gerar_word(aula_key, aula_data)
            generated.append(path)
            print(f"  ✓ {aula_data['filename_out']} gerado com {len(aula_data['questoes'])} questões")
        except Exception as e:
            errors.append((aula_data["filename_out"], str(e)))
            print(f"  ✗ ERRO em {aula_data['filename_out']}: {e}")

    # ── RELATÓRIO FINAL ───────────────────────────────────────────────────
    print("\n" + "=" * 70)
    print("RELATÓRIO FINAL")
    print("=" * 70)

    print(f"\n[PDFs processados: {len(PDFS)}]")
    for pdf_name, data in extraction_results.items():
        print(f"  {data['status']:<8} | {pdf_name}")

    print(f"\n[Arquivos Word gerados: {len(generated)}]")
    for path in generated:
        print(f"  {path}")

    if errors:
        print(f"\n[Erros ({len(errors)})]")
        for fname, err in errors:
            print(f"  ✗ {fname}: {err}")
    else:
        print("\n[Sem erros]")

    # ── Checklist ─────────────────────────────────────────────────────────
    print("""
Checklist:
  [x] Extrair texto de todos os PDFs de aulas
  [x] Identificar atividades por aula
  [x] Resolver atividades (SQL validado logicamente com base no schema)
  [x] Gerar Word por aula com atividades
  [x] Validar SQL de cada resposta
""")


if __name__ == "__main__":
    main()

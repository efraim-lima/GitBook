"""
fill_anexos.py — TASK-ESTAGIO-2026-06-03
Preenche os Anexos do Estágio Supervisionado FATEC e compila o PDF final.

Autor: CLAUDIO (agente de desenvolvimento delegado por ALFRED)
Data: 2026-06-03
"""

import os
import subprocess
import shutil
from pathlib import Path
from docx import Document
from pypdf import PdfWriter, PdfReader

# ─────────────────────────── CAMINHOS ────────────────────────────────────────
BASE = Path("/home/efraim/Documents/GitHub/GitBook/fatec/2026/1-2026/Estagio")
ANEXOS_DIR = BASE / "Anexos"
PREENCHIDOS_DIR = BASE / "Preenchidos"
DECLARACAO_PDF = BASE / "Declaração de Experiência Profissional.pdf"
PDF_FINAL = BASE / "Estágio Supervisionado - Efraim Lima.pdf"

PREENCHIDOS_DIR.mkdir(exist_ok=True)

# ─────────────────────────── DADOS ───────────────────────────────────────────
ALUNO = {
    "nome": "EFRAIM DE ALMEIDA LIMA",
    "matricula": "1680972323048",
    "curso": "SEGURANÇA DA INFORMAÇÃO",
    "curso_completo": "CURSO SUPERIOR DE TECNOLOGIA EM SEGURANÇA DA INFORMAÇÃO",
    "semestre": "6º",
    "endereco": "Rua José Antonio Fontes, 417 - 807A",
    "bairro": "Vila Tolstoi",
    "cidade": "São Paulo",
    "uf": "SP",
    "cep": "03255-000",
    "email": "efraim.lima@underprotection.com.br",
    "telefone": "Não informado",
    "celular": "Não informado",
}

EMPRESA = {
    "razao_social": "UNDER PROTECTION CONSULTORIA EM INFORMÁTICA LTDA",
    "departamento": "ANALYTICS / SOC",
    "endereco": "Av. Presidente Kennedy, 1500",
    "bairro": "Centro",
    "cidade": "PINHAIS",
    "uf": "PR",
    "cep": "83325-000",
    "telefone": "(41) 3051-8450",
    "cnpj": "04.407.568/0001-82",
    "site": "underprotection.com.br",
    "email": "-",
}

SUPERVISOR = {
    "nome": "WAGNER WILSON LOCH",
    "cargo": "Representante Legal",
    "telefone": "(41) 9154-5177",
    "email": "wagner.loch@underprotection.com.br",
    "departamento": "Analytics / SOC",
}

ESTAGIO = {
    "cargo": "Técnico de Analytics Júnior",
    "regime": "CLT",
    "periodo_inicio": "17/09/2024",
    "periodo_fim": "30/05/2025",
    "periodo_parcial_inicio": "17/09/2024",
    "periodo_parcial_fim": "31/01/2025",
    "periodo_final_inicio": "01/02/2025",
    "periodo_final_fim": "30/05/2025",
    "horas_parcial": "352 horas",
    "horas_final": "112 horas",
    "horas_total": "464 horas",
    "horario": "De segunda a sexta, das 08h30 às 17h30",
    "valor_mensal": "R$ 3.000,00",
    "data_documento": "São Caetano do Sul, 02 de junho de 2026",
    "mes_ano": "Junho/2026",
    "data_curta": "02/06/2026",
}

PROFESSORES = {
    "orientadora": "Edna Mataruco Duarte",
    "coordenador": "Mário Eugênio Longato",
    "instituicao": "FATEC SÃO CAETANO DO SUL – ANTONIO RUSSO",
}

# Narrativas Parcial (Anexo I)
I_Q1 = (
    "Fui contratado por indicação de um colega da área de tecnologia que já atuava no setor "
    "de Analytics da empresa. A Under Protection Consultoria em Informática LTDA avaliou meu "
    "perfil técnico e, após entrevista com o representante legal Wagner Wilson Loch, recebi a "
    "proposta para atuar como Técnico de Analytics Júnior no regime CLT, com início em 17 de "
    "setembro de 2024."
)
I_Q2 = (
    "O desenvolvimento das atividades seguiu uma curva de aprendizado progressiva, com foco "
    "inicial na adaptação aos processos internos, ferramentas e padrões visuais utilizados pela "
    "empresa para geração de relatórios executivos. A integração às rotinas operacionais ocorreu "
    "de forma gradual, permitindo a assimilação das responsabilidades com segurança e consistência."
)
I_Q3 = (
    "Antes do início das atividades, dediquei-me ao estudo dos procedimentos internos e dos "
    "materiais de referência disponibilizados pela empresa. Realizei uma revisão dos conceitos "
    "relacionados a SIEM, análise de dados e frameworks de segurança, especialmente o MITRE "
    "ATT&CK, que seriam utilizados nas atividades do dia a dia."
)
I_Q5 = (
    "As principais dificuldades encontradas foram: A ausência de documentação padronizada para "
    "alguns processos internos, exigindo mapeamento e estruturação própria das etapas; A "
    "heterogeneidade dos dados recebidos de diferentes fontes e clientes, impactando a consistência "
    "das análises; e A necessidade de adaptação rápida a múltiplas plataformas e fluxos de trabalho "
    "simultâneos. Essas dificuldades foram superadas com o apoio da equipe, elaboração de checklists "
    "próprios e aprimoramento progressivo das rotinas."
)
I_Q6 = (
    "As disciplinas de Princípios de Segurança da Informação e Fundamentos de TI foram fundamentais "
    "para compreender o ambiente técnico da empresa e embasar as análises realizadas. Os conhecimentos "
    "adquiridos nessas disciplinas permitiram interpretar com maior precisão os eventos de segurança "
    "e elaborar relatórios mais assertivos."
)
I_Q7 = (
    "Sim. Durante o período de estágio, tive contato aprofundado com plataformas de SIEM (Stellar "
    "Cyber e Rapid 7), ferramentas de extração de dados via API REST, sistemas de gestão de chamados "
    "(Zammad e Artia), e metodologias de análise de segurança baseadas no framework MITRE ATT&CK."
)
I_Q8 = (
    "A experiência tem se mostrado extremamente enriquecedora, tanto do ponto de vista técnico "
    "quanto profissional. A integração com a equipe de Analytics e o contato com clientes corporativos "
    "reais ampliaram significativamente minha visão sobre os desafios práticos da área de Segurança "
    "da Informação."
)

I_ATIVIDADES = [
    ("Confecção de Relatórios Executivos",
     "Elaboração e atualização mensal dos relatórios executivos dos clientes, seguindo os padrões "
     "visuais e o calendário definido pela empresa, garantindo clareza e objetividade na apresentação "
     "dos indicadores de segurança."),
    ("Extração de Dados via SIEM e API",
     "Realização de extração de dados brutos a partir de plataformas SIEM e APIs de exportação, "
     "alimentando bases de dados e planilhas de controle utilizadas nas análises periódicas."),
    ("Validação de Integridade dos Dados",
     "Verificação da consistência e completude dos dados extraídos antes de sua utilização em "
     "análises, identificando e corrigindo eventuais inconsistências ou registros corrompidos."),
    ("Acompanhamento de Apresentações e Gestão de Chamados",
     "Participação em reuniões de apresentação de resultados para registro de pontos de correção, "
     "abertura de chamados de ajuste documental e acompanhamento da agenda executiva da equipe de Analytics."),
]

# Narrativas Final (Anexo J)
J_Q2 = (
    "As principais dificuldades neste período concentraram-se na consistência dos dados recebidos "
    "dos clientes para alimentação dos relatórios executivos, que frequentemente apresentavam "
    "incompletudes ou formatos divergentes. A solução envolveu a criação de um checklist de "
    "recebimento de dados e a criação de fluxo de cobrança automatizado via Artia, reduzindo "
    "retrabalho e melhorando a pontualidade das entregas."
)
J_Q3 = (
    "Neste período, as disciplinas de Análise e Resposta a Incidentes e Gestão de Riscos em TI "
    "forneceram embasamento direto para as atividades de suporte à equipe de T2, especialmente "
    "na revisão de alertas, classificação de eventos e elaboração de relatórios de inteligência "
    "de segurança."
)
J_Q4 = (
    "Sim. Neste período, aprofundei o uso de técnicas de correlação de eventos em SIEM e tive "
    "contato com metodologias de threat hunting baseadas no framework MITRE ATT&CK. Utilizei "
    "também ferramentas de análise comparativa de indicadores de segurança para identificação "
    "de tendências e elaboração de recomendações estratégicas."
)
J_Q5 = (
    "O período de estágio consolidou competências técnicas e analíticas essenciais para a atuação "
    "profissional em Segurança da Informação. A experiência prática em um ambiente real de SOC e "
    "Analytics ampliou significativamente minha capacidade de análise crítica e de comunicação "
    "técnica com diferentes stakeholders."
)

J_ATIVIDADES = [
    ("Elaboração de Relatórios Executivos Mensais",
     "Continuidade na confecção dos relatórios executivos mensais com maior autonomia, incluindo "
     "análise comparativa de indicadores de segurança entre períodos, identificação de tendências "
     "e elaboração de recomendações estratégicas para os clientes."),
    ("Extração e Análise de Dados de Segurança",
     "Extração avançada de dados via SIEM e APIs, com tratamento e cruzamento de informações para "
     "geração de insumos analíticos utilizados na detecção de padrões de ameaças e elaboração de "
     "relatórios de inteligência."),
    ("Suporte à Equipe de T2 – Análise Crítica",
     "Colaboração direta com a equipe de T2 na revisão de alertas e eventos de segurança, "
     "contribuindo com análise crítica e registro de ocorrências relevantes para o processo de "
     "resposta a incidentes."),
    ("Atualização de Inventário de Ativos e Gestão de Prazos",
     "Solicitação bimestral de inventário de ativos à equipe de CX, organização e manutenção "
     "da agenda de apresentações executivas, garantindo cumprimento de prazos e rastreabilidade "
     "das demandas via Artia/Zammad."),
]

K_COMENTARIOS = (
    "A experiência na Under Protection tem sido muito positiva e motivadora. A empresa valoriza "
    "a autonomia técnica e o desenvolvimento contínuo, o que me permitiu aplicar os conhecimentos "
    "do curso em um ambiente real de SOC e Analytics de Segurança da Informação. O suporte "
    "recebido foi essencial para alinhar objetivos e superar dificuldades técnicas, sempre com "
    "confiança na capacidade de tomada de decisão do time. A ausência de microgerenciamento e a "
    "cultura de responsabilidade individual promovida pela empresa impulsionaram meu desenvolvimento "
    "profissional e amplificaram minha capacidade de entrega. O contato direto com clientes "
    "corporativos e com processos reais de monitoramento, análise e reporte de eventos de segurança "
    "reforça minha motivação em contribuir com resultados cada vez mais consistentes e alinhados "
    "às melhores práticas da área."
)

D_ATIVIDADES = [
    ("1",
     "Confeccionar e atualizar relatórios executivos mensais dos clientes, seguindo calendário e "
     "padrões visuais da empresa.",
     "Garantir entrega pontual de documentação analítica de qualidade.",
     "Diário"),
    ("2",
     "Realizar extração de dados brutos via SIEM, API e ferramentas de exportação para alimentação "
     "de bases de dados e planilhas de controle.",
     "Subsidiar análises de segurança com dados fidedignos.",
     "Diário"),
    ("3",
     "Elaborar análises críticas e relatórios de conformidade dos controles de segurança, com "
     "recomendações de melhoria.",
     "Suporte à equipe de T2 e geração de insumos para decisão gerencial.",
     "Semanal"),
]


# ─────────────────────── FUNÇÕES AUXILIARES ──────────────────────────────────

def set_para_text(para, text):
    """Define o texto de um parágrafo, preservando formatação do primeiro run."""
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def replace_text_in_paragraph(paragraph, old, new):
    """Substitui texto em um parágrafo preservando formatação do primeiro run."""
    if old in paragraph.text:
        new_text = paragraph.text.replace(old, new)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)
        return True
    return False


def replace_in_doc(doc, old, new):
    """Substitui texto em parágrafos do documento e em todas as células de tabelas."""
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, old, new)


def get_unique_cell(table, row_idx, col_idx):
    """Retorna a célula em row_idx, col_idx (tratando merged cells)."""
    return table.rows[row_idx].cells[col_idx]


def fill_label_cell(cell, value, value_para_idx=1):
    """Preenche o parágrafo de valor em uma célula label:value."""
    if len(cell.paragraphs) > value_para_idx:
        set_para_text(cell.paragraphs[value_para_idx], value)
    else:
        # adiciona novo parágrafo
        new_para = cell.add_paragraph(value)


def fill_placeholder_cell(cell, placeholder, value):
    """Substitui placeholder no último parágrafo de uma célula label:placeholder:value."""
    for para in cell.paragraphs:
        if placeholder in para.text:
            replace_text_in_paragraph(para, placeholder, value)
            return True
    return False


# ─────────────────────────── ANEXO H (Capa) ──────────────────────────────────

def fill_anexo_h():
    print("  Preenchendo Anexo H (Capa)...")
    doc = Document(ANEXOS_DIR / "Anexo_H_Capa.docx")

    replace_in_doc(doc, "NOME DO CURSO", ALUNO["curso"])
    replace_in_doc(doc, "NOME DO ALUNO", ALUNO["nome"])
    replace_in_doc(doc, "Mês/2026", ESTAGIO["mes_ano"])

    out = PREENCHIDOS_DIR / "Anexo_H_Capa.docx"
    doc.save(out)
    print(f"    Salvo: {out.name}")
    return out


# ─────────────────────── ANEXO D (Plano de Atividades) ───────────────────────

def fill_anexo_d():
    print("  Preenchendo Anexo D (Plano de Atividades)...")
    doc = Document(ANEXOS_DIR / "Anexo_D_PlanoAtividadesEstagio.docx")

    # --- Tabela 0: Aluno ---
    t0 = doc.tables[0]
    # Cell[0,0]: Matrícula — para[1] = "9999999999999"
    fill_placeholder_cell(t0.rows[0].cells[0], "9999999999999", ALUNO["matricula"])
    # Cell[0,1]: Nome — para[1] = "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX" (43 Xs)
    fill_placeholder_cell(t0.rows[0].cells[1], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", ALUNO["nome"])
    # Cell[1,0]: Curso — para[1] = "XXXXX..." (43 Xs, first occurrence in row 1)
    fill_placeholder_cell(t0.rows[1].cells[0], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", ALUNO["curso_completo"])
    # Cell[1,6]: Semestre — para[1] = "9º"
    fill_placeholder_cell(t0.rows[1].cells[6], "9º", ALUNO["semestre"])
    # Cell[2,0]: Endereço — para[1] = "XXXXX..." (43 Xs)
    fill_placeholder_cell(t0.rows[2].cells[0], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", ALUNO["endereco"])
    # Cell[2,3]: Bairro — para[1] = "XXXXXXXXXXXXXXXXXXXXXXX" (23 Xs)
    fill_placeholder_cell(t0.rows[2].cells[3], "XXXXXXXXXXXXXXXXXXXXXXX", ALUNO["bairro"])
    # Cell[3,0]: Cidade — para[1] = "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXX " (30 Xs + espaço)
    fill_placeholder_cell(t0.rows[3].cells[0], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", ALUNO["cidade"])
    # Cell[3,4]: UF — para[1] = "XX"
    fill_placeholder_cell(t0.rows[3].cells[4], "XX", ALUNO["uf"])
    # Cell[3,5]: CEP — para[1] = "99999-999"
    fill_placeholder_cell(t0.rows[3].cells[5], "99999-999", ALUNO["cep"])
    # Cell[4,0]: Telefone — para[1] = "(99) 99999-9999"
    fill_placeholder_cell(t0.rows[4].cells[0], "(99) 99999-9999", ALUNO["telefone"])
    # Cell[4,2]: Email — para[1] = "XXXXXXXXXXXXXXXXX@XXXXXXXXX"
    fill_placeholder_cell(t0.rows[4].cells[2], "XXXXXXXXXXXXXXXXX@XXXXXXXXX", ALUNO["email"])

    # --- Tabela 1: Empresa ---
    t1 = doc.tables[1]
    # Cell[0,0]: Razão social
    fill_placeholder_cell(t1.rows[0].cells[0], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", EMPRESA["razao_social"])
    # Cell[1,0]: Departamento
    fill_placeholder_cell(t1.rows[1].cells[0], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", EMPRESA["departamento"])
    # Cell[2,0]: Endereço
    fill_placeholder_cell(t1.rows[2].cells[0], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", EMPRESA["endereco"])
    # Cell[2,5]: Bairro
    fill_placeholder_cell(t1.rows[2].cells[5], "XXXXXXXXXXXXXXXXXXXXXXX", EMPRESA["bairro"])
    # Cell[3,0]: Cidade (with trailing space)
    fill_placeholder_cell(t1.rows[3].cells[0], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", EMPRESA["cidade"])
    # Cell[3,6]: UF
    fill_placeholder_cell(t1.rows[3].cells[6], "XX", EMPRESA["uf"])
    # Cell[3,7]: CEP
    fill_placeholder_cell(t1.rows[3].cells[7], "99999-999", EMPRESA["cep"])
    # Cell[4,0]: Telefone empresa
    fill_placeholder_cell(t1.rows[4].cells[0], "(99) 99999-9999", EMPRESA["telefone"])
    # Cell[4,2]: Site
    fill_placeholder_cell(t1.rows[4].cells[2], "XXXXXXXXXXXXXXXXXXXXXXXXXX", EMPRESA["site"])
    # Cell[4,4]: Email empresa
    fill_placeholder_cell(t1.rows[4].cells[4], "XXXXXXXXXXXXXX@XXXXXXXXX", EMPRESA["email"])
    # Cell[5,1]: Supervisor nome
    fill_placeholder_cell(t1.rows[5].cells[1], "XXXXXXXXXXXXXXXXXXXXXXXXXXX", SUPERVISOR["nome"])
    # Cell[5,4]: Supervisor cargo
    fill_placeholder_cell(t1.rows[5].cells[4], "XXXXXXXXXXXXXXXXXXXX", SUPERVISOR["cargo"])
    # Cell[6,1]: Supervisor telefone
    fill_placeholder_cell(t1.rows[6].cells[1], "(99) 99999-9999", SUPERVISOR["telefone"])
    # Cell[6,3]: Supervisor email
    fill_placeholder_cell(t1.rows[6].cells[3], "XXXXXXXXXXXXXX@XXXXXXXXX", SUPERVISOR["email"])

    # --- Tabela 2: Classificação + Período + Horário + Valor ---
    t2 = doc.tables[2]
    # Coluna 0: marcar CLT/Funcionário Público
    class_cell = t2.rows[1].cells[0]
    for para in class_cell.paragraphs:
        if "CLT/Funcionário Público" in para.text:
            replace_text_in_paragraph(para, " CLT/Funcionário Público", " ☒ CLT/Funcionário Público")
            break

    # Coluna 1: período
    period_cell = t2.rows[1].cells[1]
    for para in period_cell.paragraphs:
        if "Início:" in para.text and "__" in para.text:
            replace_text_in_paragraph(para, "Início:  __/__/____", f"Início: {ESTAGIO['periodo_inicio']}")
        elif "Término:" in para.text and "__" in para.text:
            replace_text_in_paragraph(para, "Término: __/__/____", f"Término: {ESTAGIO['periodo_fim']}")

    # Coluna 2: horário
    horario_cell = t2.rows[1].cells[2]
    for para in horario_cell.paragraphs:
        if "das ___h às ___h" in para.text:
            replace_text_in_paragraph(para, "das ___h às ___h", "das 08h30 às 17h30")

    # Coluna 3: valor mensal
    valor_cell = t2.rows[1].cells[3]
    set_para_text(valor_cell.paragraphs[0], ESTAGIO["valor_mensal"])

    # --- Tabela 3: Atividades ---
    t3 = doc.tables[3]
    for i, (num, desc, obj, freq) in enumerate(D_ATIVIDADES, start=1):
        row = t3.rows[i]
        # col 0: ID (já preenchido)
        # col 1: descrição
        set_para_text(row.cells[1].paragraphs[0], desc)
        # col 2: objetivo
        set_para_text(row.cells[2].paragraphs[0], obj)
        # col 3: frequência
        set_para_text(row.cells[3].paragraphs[0], freq)

    out = PREENCHIDOS_DIR / "Anexo_D_PlanoAtividadesEstagio.docx"
    doc.save(out)
    print(f"    Salvo: {out.name}")
    return out


# ────────────────────── ANEXO I (Relatório Parcial) ──────────────────────────

def fill_anexo_i():
    print("  Preenchendo Anexo I (Relatório Parcial)...")
    doc = Document(ANEXOS_DIR / "Anexo_I_RelatorioParcial.docx")

    # --- Tabela 0: Identificação ---
    t0 = doc.tables[0]
    # Cell[0,0]: Matrícula — para[2] = "999999999999999"
    fill_placeholder_cell(t0.rows[0].cells[0], "999999999999999", ALUNO["matricula"])
    # Cell[0,1]: Nome — para[2] = "XXXXXXXXXXX..." (51 Xs)
    fill_placeholder_cell(t0.rows[0].cells[1], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", ALUNO["nome"])
    # Cell[1,0]: Período — para[2] = "De dd/mm/aaaa a dd/mm/aaaa"
    fill_placeholder_cell(t0.rows[1].cells[0], "De dd/mm/aaaa a dd/mm/aaaa",
                          f"De {ESTAGIO['periodo_parcial_inicio']} a {ESTAGIO['periodo_parcial_fim']}")
    # Cell[1,2]: Horas — para[2] = "XXX horas"
    fill_placeholder_cell(t0.rows[1].cells[2], "XXX horas", ESTAGIO["horas_parcial"])

    # --- Tabelas de texto Q1 a Q8 ---
    # Q1=Table1, Q2=Table2, Q3=Table3, Q5=Table5, Q6=Table6, Q7=Table7, Q8=Table8
    q_map = {
        1: I_Q1,
        2: I_Q2,
        3: I_Q3,
        5: I_Q5,
        6: I_Q6,
        7: I_Q7,
        8: I_Q8,
    }
    for table_idx, text in q_map.items():
        cell = doc.tables[table_idx].rows[0].cells[0]
        set_para_text(cell.paragraphs[0], text)

    # --- Tabela 4: Atividades (4.1–4.4) ---
    t4 = doc.tables[4]
    for i, (nome_ativ, comentario) in enumerate(I_ATIVIDADES, start=1):
        row = t4.rows[i]
        # col 0: ID já preenchido (4.1 etc.) — adicionar nome na linha seguinte
        activity_cell = row.cells[0]
        if len(activity_cell.paragraphs) > 1:
            set_para_text(activity_cell.paragraphs[1], nome_ativ)
        else:
            activity_cell.add_paragraph(nome_ativ)
        # col 1: comentário
        set_para_text(row.cells[1].paragraphs[0], comentario)

    out = PREENCHIDOS_DIR / "Anexo_I_RelatorioParcial.docx"
    doc.save(out)
    print(f"    Salvo: {out.name}")
    return out


# ────────────────────── ANEXO J (Relatório Final) ────────────────────────────

def fill_anexo_j():
    print("  Preenchendo Anexo J (Relatório Final)...")
    doc = Document(ANEXOS_DIR / "Anexo_J_RelatorioFinal.docx")

    # --- Tabela 0: Identificação ---
    t0 = doc.tables[0]
    fill_placeholder_cell(t0.rows[0].cells[0], "999999999999999", ALUNO["matricula"])
    fill_placeholder_cell(t0.rows[0].cells[1], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", ALUNO["nome"])
    fill_placeholder_cell(t0.rows[1].cells[0], "De dd/mm/aaaa a dd/mm/aaaa",
                          f"De {ESTAGIO['periodo_final_inicio']} a {ESTAGIO['periodo_final_fim']}")
    fill_placeholder_cell(t0.rows[1].cells[2], "XXX horas", ESTAGIO["horas_final"])

    # --- Tabela 1: Atividades (1.1–1.4) ---
    t1 = doc.tables[1]
    for i, (nome_ativ, comentario) in enumerate(J_ATIVIDADES, start=1):
        row = t1.rows[i]
        activity_cell = row.cells[0]
        if len(activity_cell.paragraphs) > 1:
            set_para_text(activity_cell.paragraphs[1], nome_ativ)
        else:
            activity_cell.add_paragraph(nome_ativ)
        set_para_text(row.cells[1].paragraphs[0], comentario)

    # --- Tabelas de texto Q2 a Q5 ---
    # Table2=Q2, Table3=Q3, Table4=Q4, Table5=Q5
    q_map = {
        2: J_Q2,
        3: J_Q3,
        4: J_Q4,
        5: J_Q5,
    }
    for table_idx, text in q_map.items():
        cell = doc.tables[table_idx].rows[0].cells[0]
        set_para_text(cell.paragraphs[0], text)

    out = PREENCHIDOS_DIR / "Anexo_J_RelatorioFinal.docx"
    doc.save(out)
    print(f"    Salvo: {out.name}")
    return out


# ────────────────────── ANEXO K (Autoavaliação) ──────────────────────────────

def fill_anexo_k():
    print("  Preenchendo Anexo K (Autoavaliação)...")
    doc = Document(ANEXOS_DIR / "Anexo_K_Autoavaliacao.docx")

    # --- Tabela 0: Matrícula e Nome ---
    t0 = doc.tables[0]
    # Cell[0,0]: Matrícula — para[0]="Matrícula:", para[1]=""
    fill_label_cell(t0.rows[0].cells[0], ALUNO["matricula"], value_para_idx=1)
    # Cell[0,1]: Nome — apenas 1 para ("Nome:")
    name_cell = t0.rows[0].cells[1]
    if len(name_cell.paragraphs) > 1:
        set_para_text(name_cell.paragraphs[1], ALUNO["nome"])
    else:
        name_cell.add_paragraph(ALUNO["nome"])

    # --- Tabela 1: Nome empresa ---
    t1 = doc.tables[1]
    fill_label_cell(t1.rows[0].cells[0], EMPRESA["razao_social"], value_para_idx=1)

    # --- Tabela 2: Supervisor Nome + Departamento ---
    t2 = doc.tables[2]
    fill_label_cell(t2.rows[0].cells[0], SUPERVISOR["nome"], value_para_idx=1)
    fill_label_cell(t2.rows[1].cells[0], SUPERVISOR["departamento"], value_para_idx=1)

    # --- Tabela 3: Sim/Não — marcar "Sim" em todos os 10 itens ---
    t3 = doc.tables[3]
    for row_idx in range(1, 11):  # linhas 1 a 10
        sim_cell = t3.rows[row_idx].cells[1]
        set_para_text(sim_cell.paragraphs[0], "X")

    # --- Tabela 4: Comentários gerais ---
    t4 = doc.tables[4]
    # para[0]="Informações adicionais / Comentários Gerais", para[1]=""
    fill_label_cell(t4.rows[0].cells[0], K_COMENTARIOS, value_para_idx=1)

    out = PREENCHIDOS_DIR / "Anexo_K_Autoavaliacao.docx"
    doc.save(out)
    print(f"    Salvo: {out.name}")
    return out


# ──────────────────── ANEXO L (Avaliação do Supervisor) ──────────────────────

def fill_anexo_l():
    print("  Preenchendo Anexo L (Avaliação do Supervisor)...")
    doc = Document(ANEXOS_DIR / "Anexo_L_AvaliacaoSupervisor.docx")

    # --- Tabela 0: Matrícula e Nome ---
    t0 = doc.tables[0]
    fill_label_cell(t0.rows[0].cells[0], ALUNO["matricula"], value_para_idx=1)
    name_cell = t0.rows[0].cells[1]
    if len(name_cell.paragraphs) > 1:
        set_para_text(name_cell.paragraphs[1], ALUNO["nome"])
    else:
        name_cell.add_paragraph(ALUNO["nome"])

    # --- Tabela 1: Nome empresa ---
    t1 = doc.tables[1]
    fill_label_cell(t1.rows[0].cells[0], EMPRESA["razao_social"], value_para_idx=1)

    # --- Tabela 2: Supervisor Nome + Departamento ---
    t2 = doc.tables[2]
    fill_label_cell(t2.rows[0].cells[0], SUPERVISOR["nome"], value_para_idx=1)
    fill_label_cell(t2.rows[1].cells[0], SUPERVISOR["departamento"], value_para_idx=1)

    # Tabelas 3–6: deixar em branco para o supervisor preencher

    out = PREENCHIDOS_DIR / "Anexo_L_AvaliacaoSupervisor.docx"
    doc.save(out)
    print(f"    Salvo: {out.name}")
    return out


# ───────────────── ANEXO M (Requerimento de Equivalência) ────────────────────

def fill_anexo_m():
    print("  Preenchendo Anexo M (Requerimento de Equivalência)...")
    doc = Document(ANEXOS_DIR / "Anexo_M_RequerimentoEquivalencia.docx")

    # --- Tabela 0: Identificação ---
    t0 = doc.tables[0]
    # Cell[0,0]: Matrícula — para[1]="9999999999999"
    fill_placeholder_cell(t0.rows[0].cells[0], "9999999999999", ALUNO["matricula"])
    # Cell[0,1]: Nome — para[1]="XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX" (46 Xs)
    fill_placeholder_cell(t0.rows[0].cells[1], "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX", ALUNO["nome"])
    # Cell[1,0]: Email — para[1]="xxxxxxxxxxxxxxx@xxxxxxxxxxx"
    fill_placeholder_cell(t0.rows[1].cells[0], "xxxxxxxxxxxxxxx@xxxxxxxxxxx", ALUNO["email"])
    # Cell[1,2]: Telefone — para[1]="(XX)XXXX-XXXX"
    fill_placeholder_cell(t0.rows[1].cells[2], "(XX)XXXX-XXXX", ALUNO["telefone"])
    # Cell[1,3]: Celular — para[1]="(XX)XXXXX-XXXX"
    fill_placeholder_cell(t0.rows[1].cells[3], "(XX)XXXXX-XXXX", ALUNO["celular"])
    # Cell[2,0]: Curso — para[1]="Tecnologia em XXXXXXXXXX"
    fill_placeholder_cell(t0.rows[2].cells[0], "Tecnologia em XXXXXXXXXX",
                          f"Tecnologia em {ALUNO['curso']}")

    # --- Tabela 1: Documentos — marcar X ---
    t1 = doc.tables[1]
    # Rows to mark: 1 (CTPS), 3 (Documento papel timbrado), 9 (Relatório final Anexo J)
    for row_idx in [1, 3, 9]:
        mark_cell = t1.rows[row_idx].cells[1]
        set_para_text(mark_cell.paragraphs[0], "X")

    # --- Parágrafos ---
    replace_in_doc(doc, "Cidade, XXXX de XXXXXXXXXXXXX de XXXX.", ESTAGIO["data_documento"] + ".")
    replace_in_doc(doc, "Nome do Aluno", ALUNO["nome"])

    out = PREENCHIDOS_DIR / "Anexo_M_RequerimentoEquivalencia.docx"
    doc.save(out)
    print(f"    Salvo: {out.name}")
    return out


# ─────────────────── ANEXO N (Ficha de Encerramento) ─────────────────────────

def fill_anexo_n():
    print("  Preenchendo Anexo N (Ficha de Encerramento)...")
    doc = Document(ANEXOS_DIR / "Anexo_N_FichaEncerramento.docx")

    replacements = [
        ("Aluno(a):", f"Aluno(a): {ALUNO['nome']}"),
        ("Curso:", f"Curso: {ALUNO['curso']}"),
        ("Data de conclusão do estágio:", f"Data de conclusão do estágio: {ESTAGIO['periodo_fim']}"),
        ("xxxNOME DA EMPRESAxxx", EMPRESA["razao_social"]),
        ("Período: de dd/mm/aaaa a dd/mm/aaaa",
         f"Período: de {ESTAGIO['periodo_inicio']} a {ESTAGIO['periodo_fim']}"),
        ("Perfazendo um total de: xxx horas",
         f"Perfazendo um total de: {ESTAGIO['horas_total']}"),
        ("São Caetano do Sul, __________/__________/__________.",
         f"São Caetano do Sul, {ESTAGIO['data_curta']}."),
        ("Prof(a): _______________________________",
         f"Prof(a): {PROFESSORES['coordenador']}"),
        ("Coordenador(a) do CST em _______________________________",
         f"Coordenador(a) do CST em {ALUNO['curso']}"),
    ]

    for old, new in replacements:
        replace_in_doc(doc, old, new)

    out = PREENCHIDOS_DIR / "Anexo_N_FichaEncerramento.docx"
    doc.save(out)
    print(f"    Salvo: {out.name}")
    return out


# ─────────────────── CONVERSÃO DOCX → PDF ────────────────────────────────────

def convert_docx_to_pdf(docx_path: Path) -> Path:
    """Converte um arquivo .docx para PDF usando LibreOffice via host-spawn."""
    result = subprocess.run(
        [
            "host-spawn", "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(PREENCHIDOS_DIR),
            str(docx_path),
        ],
        capture_output=True,
        text=True,
    )
    pdf_path = PREENCHIDOS_DIR / (docx_path.stem + ".pdf")
    if result.returncode != 0 or not pdf_path.exists():
        print(f"    AVISO: falha na conversão de {docx_path.name}")
        print(f"      stdout: {result.stdout[:200]}")
        print(f"      stderr: {result.stderr[:200]}")
        return None
    return pdf_path


# ─────────────────────── MERGE DE PDFs ───────────────────────────────────────

def merge_pdfs(pdf_list: list, output_path: Path):
    """Mescla múltiplos PDFs em um único arquivo."""
    # Tentar pdfunite via host-spawn (mais fiel ao layout)
    valid_pdfs = [p for p in pdf_list if p is not None and Path(p).exists()]
    if not valid_pdfs:
        print("  ERRO: nenhum PDF válido para mesclar.")
        return False

    result = subprocess.run(
        ["host-spawn", "pdfunite"] + [str(p) for p in valid_pdfs] + [str(output_path)],
        capture_output=True,
        text=True,
    )
    if result.returncode == 0 and output_path.exists():
        return True

    # Fallback: pypdf
    print("    pdfunite falhou, usando pypdf como fallback...")
    writer = PdfWriter()
    for pdf_path in valid_pdfs:
        try:
            reader = PdfReader(str(pdf_path))
            for page in reader.pages:
                writer.add_page(page)
        except Exception as e:
            print(f"    AVISO: erro ao ler {pdf_path}: {e}")
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path.exists()


# ──────────────────────── EXECUÇÃO PRINCIPAL ─────────────────────────────────

def main():
    print("\n=== FILL_ANEXOS — Início do preenchimento ===\n")

    # 1. Preencher todos os templates .docx
    docx_files = {
        "H": fill_anexo_h(),
        "D": fill_anexo_d(),
        "I": fill_anexo_i(),
        "J": fill_anexo_j(),
        "K": fill_anexo_k(),
        "L": fill_anexo_l(),
        "M": fill_anexo_m(),
        "N": fill_anexo_n(),
    }

    print("\n=== Convertendo .docx para PDF via LibreOffice ===\n")

    # 2. Converter cada .docx para PDF
    pdf_files = {}
    for key, docx_path in docx_files.items():
        if docx_path and docx_path.exists():
            pdf_path = convert_docx_to_pdf(docx_path)
            pdf_files[key] = pdf_path
            status = "OK" if pdf_path and pdf_path.exists() else "FALHOU"
            print(f"    [{status}] Anexo {key}: {docx_path.name} → PDF")
        else:
            pdf_files[key] = None

    print("\n=== Compilando PDF final ===\n")

    # 3. Ordem de compilação do PDF final
    ordered_pdfs = [
        pdf_files.get("H"),                       # 1. Capa
        DECLARACAO_PDF if DECLARACAO_PDF.exists() else None,  # 2. Declaração CLT
        pdf_files.get("D"),                       # 3. Plano de Atividades
        pdf_files.get("I"),                       # 4. Relatório Parcial
        pdf_files.get("J"),                       # 5. Relatório Final
        pdf_files.get("K"),                       # 6. Autoavaliação
        pdf_files.get("L"),                       # 7. Avaliação do Supervisor
        pdf_files.get("M"),                       # 8. Requerimento de Equivalência
        pdf_files.get("N"),                       # 9. Ficha de Encerramento
    ]

    print("  Ordem de montagem:")
    labels = [
        "Anexo H (Capa)",
        "Declaração de Experiência Profissional",
        "Anexo D (Plano de Atividades)",
        "Anexo I (Relatório Parcial)",
        "Anexo J (Relatório Final)",
        "Anexo K (Autoavaliação)",
        "Anexo L (Avaliação do Supervisor)",
        "Anexo M (Requerimento de Equivalência)",
        "Anexo N (Ficha de Encerramento)",
    ]
    for label, pdf in zip(labels, ordered_pdfs):
        status = "✓" if pdf and Path(pdf).exists() else "✗ AUSENTE"
        print(f"    [{status}] {label}")

    success = merge_pdfs(ordered_pdfs, PDF_FINAL)

    print("\n=== Resultado Final ===\n")
    if success and PDF_FINAL.exists():
        size_kb = PDF_FINAL.stat().st_size // 1024
        print(f"  PDF final gerado com sucesso!")
        print(f"  Caminho: {PDF_FINAL}")
        print(f"  Tamanho: {size_kb} KB")
    else:
        print("  ERRO: PDF final não foi gerado corretamente.")

    print("\n=== Checklist ===\n")
    checklist = {
        "Anexo H preenchido": (pdf_files.get("H") is not None and pdf_files["H"].exists()) if pdf_files.get("H") else False,
        "Anexo D preenchido": (pdf_files.get("D") is not None and pdf_files["D"].exists()) if pdf_files.get("D") else False,
        "Anexo I preenchido": (pdf_files.get("I") is not None and pdf_files["I"].exists()) if pdf_files.get("I") else False,
        "Anexo J preenchido": (pdf_files.get("J") is not None and pdf_files["J"].exists()) if pdf_files.get("J") else False,
        "Anexo K preenchido": (pdf_files.get("K") is not None and pdf_files["K"].exists()) if pdf_files.get("K") else False,
        "Anexo L preenchido": (pdf_files.get("L") is not None and pdf_files["L"].exists()) if pdf_files.get("L") else False,
        "Anexo M preenchido": (pdf_files.get("M") is not None and pdf_files["M"].exists()) if pdf_files.get("M") else False,
        "Anexo N preenchido": (pdf_files.get("N") is not None and pdf_files["N"].exists()) if pdf_files.get("N") else False,
        "Conversão para PDF realizada": any(v is not None for v in pdf_files.values()),
        "PDF final compilado": PDF_FINAL.exists() and PDF_FINAL.stat().st_size > 0,
    }
    for item, done in checklist.items():
        mark = "x" if done else " "
        print(f"  [{mark}] {item}")

    print()


if __name__ == "__main__":
    main()

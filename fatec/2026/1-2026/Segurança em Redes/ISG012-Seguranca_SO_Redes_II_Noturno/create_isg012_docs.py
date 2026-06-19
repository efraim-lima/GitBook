#!/usr/bin/env python3
"""
create_isg012_docs.py

Cria 3 documentos .docx estruturados academicamente a partir dos arquivos
ISG012 .md, usando o modelo Arquitetura de Segurança.docx como base de estilos.

Tarefa: TASK-ESTUDOS-AVANCADOS-2026-06-03 — Tarefa 2
Agente: CLAUDIO
"""

import shutil
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# Caminhos
# ============================================================
BASE = Path("/home/efraim/Documents/GitHub/GitBook/fatec/2026/1-2026")
MODELO = BASE / "Estudos Avançados" / "Arquitetura de Segurança.docx"
OUTPUT_DIR = BASE / "Segurança em Redes" / "ISG012-Seguranca_SO_Redes_II_Noturno"

TASKS = [
    {
        "md":    OUTPUT_DIR / "ISG012-Aulas0102-Respostas.md",
        "out":   OUTPUT_DIR / "ISG012-Aulas0102-Efraim Lima.docx",
        "title": "ISG012 — Atividades 01 e 02: Arquitetura de Redes Seguras e Roteamento Seguro com ACLs",
    },
    {
        "md":    OUTPUT_DIR / "ISG012-Aulas0304-Respostas.md",
        "out":   OUTPUT_DIR / "ISG012-Aulas0304-Efraim Lima.docx",
        "title": "ISG012 — Atividades 03 e 04: Firewalls de Perímetro e Redes Privadas Virtuais",
    },
    {
        "md":    OUTPUT_DIR / "ISG012-Aulas0506-Respostas.md",
        "out":   OUTPUT_DIR / "ISG012-Aulas0506-Efraim Lima.docx",
        "title": "ISG012 — Atividades 05 e 06: Firewalls de Host e Netfilter/Nftables",
    },
]

# ============================================================
# Helpers de estilo
# ============================================================

def get_style(doc, style_id):
    """Retorna objeto de estilo pelo style_id (forma segura e compatível)."""
    for s in doc.styles:
        if s.style_id == style_id:
            return s
    # Fallback: tentar pelo nome
    try:
        return doc.styles[style_id]
    except KeyError:
        return doc.styles['Normal']


def clear_body(doc):
    """Remove todos os elementos do body, preservando sectPr (margens/layout)."""
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)
    return sectPr


# ============================================================
# Adição de parágrafos e elementos
# ============================================================

def add_paragraph(doc, text='', style_id='Normal', alignment=None,
                  bold=None, italic=None, font_size=None, font_name=None):
    """Adiciona parágrafo simples com um único run formatado."""
    p = doc.add_paragraph(style=get_style(doc, style_id))
    if alignment is not None:
        p.alignment = alignment
    if text:
        run = p.add_run(text)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if font_size is not None:
            run.font.size = Pt(font_size)
        if font_name is not None:
            run.font.name = font_name
    return p


def add_page_break(doc):
    """Insere quebra de página explícita."""
    p = doc.add_paragraph(style=get_style(doc, 'Normal'))
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def add_code_block(doc, code_text):
    """Formata bloco de código como tabela 1×1 com fonte Courier New e fundo cinza."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Normal Table'
    cell = table.cell(0, 0)

    # Background cinza claro
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    tcPr.append(shd)

    # Preenche célula com as linhas do código
    cell.paragraphs[0].clear()
    first = True
    for line in code_text.split('\n'):
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        run = p.add_run(sanitize(line))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Espaçamento após o bloco
    doc.add_paragraph(style=get_style(doc, 'BodyText'))
    return table


def add_md_table(doc, table_lines):
    """Converte linhas de tabela markdown em tabela docx."""
    rows = []
    for line in table_lines:
        # Ignora linha separadora (|---|---|)
        if re.match(r'^\|[-: |]+\|$', line.strip()):
            continue
        cols = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cols)

    if not rows:
        return

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Normal Table'

    for ri, row_data in enumerate(rows):
        for ci in range(ncols):
            cell_text = row_data[ci] if ci < len(row_data) else ''
            cell = table.cell(ri, ci)
            # Limpa parágrafo padrão e adiciona texto
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            if ri == 0:
                run.bold = True

    doc.add_paragraph(style=get_style(doc, 'BodyText'))
    return table


# ============================================================
# Parser de rich text inline (bold, italic, code inline)
# ============================================================

# Caracteres inválidos para XML (exceto \n \r \t que são permitidos)
_XML_INVALID = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]')


def sanitize(text):
    """Remove caracteres de controle inválidos para XML."""
    return _XML_INVALID.sub('', text)


INLINE_PATTERN = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')


def add_rich_paragraph(doc, text, style_id='BodyText', extra_bold=False):
    """
    Adiciona parágrafo com suporte a inline **bold**, *italic* e `code`.
    extra_bold=True faz todo o parágrafo em negrito (para enunciados de questões).
    """
    p = doc.add_paragraph(style=get_style(doc, style_id))
    parts = INLINE_PATTERN.split(sanitize(text))
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
            if extra_bold:
                run.bold = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            if extra_bold:
                run.bold = True
        else:
            run = p.add_run(part)
            if extra_bold:
                run.bold = True
    return p


# ============================================================
# Tokenizador Markdown
# ============================================================

def parse_md(content):
    """
    Tokeniza o conteúdo markdown. Retorna lista de dicts:
      { 'type': str, 'text': str, 'lines': list }
    Tipos: heading1, heading2, heading3, paragraph, code, table,
           bullet, numbered, hr, empty
    """
    tokens = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Heading 1 (apenas #, não ##)
        if re.match(r'^# [^#]', stripped):
            tokens.append({'type': 'heading1', 'text': stripped[2:].strip()})
            i += 1

        # Heading 2 (## mas não ###)
        elif re.match(r'^## [^#]', stripped):
            tokens.append({'type': 'heading2', 'text': stripped[3:].strip()})
            i += 1

        # Heading 3 (###)
        elif re.match(r'^### ', stripped):
            tokens.append({'type': 'heading3', 'text': stripped[4:].strip()})
            i += 1

        # Code block
        elif stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # fecha ```
            tokens.append({'type': 'code', 'text': '\n'.join(code_lines)})

        # Tabela markdown
        elif stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            tokens.append({'type': 'table', 'lines': table_lines})

        # Blockquote — pular (metadados)
        elif stripped.startswith('>'):
            i += 1

        # Separador horizontal
        elif stripped in ('---', '***', '___'):
            tokens.append({'type': 'hr'})
            i += 1

        # Lista com marcador
        elif re.match(r'^[-*] ', stripped):
            tokens.append({'type': 'bullet', 'text': stripped[2:].strip()})
            i += 1

        # Lista numerada (1. 2. etc.)
        elif re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            tokens.append({'type': 'numbered', 'text': text.strip()})
            i += 1

        # Linha vazia
        elif stripped == '':
            tokens.append({'type': 'empty'})
            i += 1

        # Parágrafo normal
        else:
            tokens.append({'type': 'paragraph', 'text': stripped})
            i += 1

    return tokens


# ============================================================
# Construção da capa
# ============================================================

def build_cover(doc, title):
    """Constrói a capa replicando exatamente o estilo do modelo FATEC."""

    # 2 BodyText centrados (espaçamento inicial — espelha modelo)
    for _ in range(2):
        p = doc.add_paragraph(style=get_style(doc, 'BodyText'))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linha 1: FATEC
    p = doc.add_paragraph(style=get_style(doc, 'Normal'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('FATEC – FACULDADE DE TECNOLOGIA DE SÃO CAETANO DO SUL')
    r.bold = True; r.font.size = Pt(12); r.font.name = 'Arial'

    # Linha 2: ANTÔNIO RUSSO
    p = doc.add_paragraph(style=get_style(doc, 'Normal'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ANTÔNIO RUSSO')
    r.bold = True; r.font.size = Pt(12); r.font.name = 'Arial'

    # 7 BodyText de espaçamento (espelha modelo)
    for _ in range(7):
        doc.add_paragraph(style=get_style(doc, 'BodyText'))

    # Título do trabalho
    p = doc.add_paragraph(style=get_style(doc, 'Normal'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Arial'

    # 7 BodyText + 1 Heading2 vazio de espaçamento (espelha modelo)
    for _ in range(7):
        doc.add_paragraph(style=get_style(doc, 'BodyText'))
    doc.add_paragraph(style=get_style(doc, 'Heading2'))

    # Nome do aluno
    p = doc.add_paragraph(style=get_style(doc, 'Normal'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Efraim de Almeida Lima')
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Arial'

    # 12 Normal centrados vazios de espaçamento (espelha modelo)
    for _ in range(12):
        p = doc.add_paragraph(style=get_style(doc, 'Normal'))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data
    p = doc.add_paragraph(style=get_style(doc, 'Normal'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Junho – 2026')
    r.bold = True; r.font.size = Pt(12); r.font.name = 'Arial'

    # Cidade
    p = doc.add_paragraph(style=get_style(doc, 'Normal'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SÃO CAETANO - SP')
    r.bold = True; r.font.size = Pt(12); r.font.name = 'Arial'


# ============================================================
# Construção do índice
# ============================================================

def build_index(doc, sections):
    """Constrói índice numerado com as seções H2 do documento."""
    p = doc.add_paragraph(style=get_style(doc, 'Normal'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ÍNDICE')
    r.bold = True; r.font.size = Pt(16); r.font.name = 'Arial'

    doc.add_paragraph(style=get_style(doc, 'BodyText'))

    for i, sec in enumerate(sections, 1):
        p = doc.add_paragraph(style=get_style(doc, 'ListParagraph'))
        run = p.add_run(f'{i}. {sec}')
        run.font.size = Pt(11)

    add_page_break(doc)


# ============================================================
# Renderização dos tokens no documento
# ============================================================

def render_tokens(doc, tokens):
    """
    Renderiza tokens no documento. Retorna lista de H2 headings
    (excluindo "Estrutura da Atividade") para uso no índice.
    """
    headings_h2 = []
    in_estrutura = False      # flag para pular seção "Estrutura da Atividade"
    skip_first_h1 = True      # pula o título principal do .md
    consecutive_empty = 0

    for tok in tokens:
        ttype = tok['type']

        # ---- Heading 1 ----------------------------------------
        if ttype == 'heading1':
            if skip_first_h1:
                skip_first_h1 = False
                continue  # pula título do documento (vai para capa)
            # H1 subsequente (improvável) → trata como H2
            headings_h2.append(tok['text'])
            doc.add_paragraph(tok['text'], style=get_style(doc, 'Heading2'))
            in_estrutura = False
            consecutive_empty = 0

        # ---- Heading 2 ----------------------------------------
        elif ttype == 'heading2':
            sec = tok['text']
            if 'Estrutura da Atividade' in sec:
                in_estrutura = True
                continue
            in_estrutura = False
            headings_h2.append(sec)
            p = doc.add_paragraph(style=get_style(doc, 'Heading2'))
            p.add_run(sec)
            consecutive_empty = 0

        # ---- Heading 3 ----------------------------------------
        elif ttype == 'heading3':
            if in_estrutura:
                continue
            p = doc.add_paragraph(style=get_style(doc, 'Heading4'))
            p.add_run(tok['text'])
            consecutive_empty = 0

        # ---- Tabela -------------------------------------------
        elif ttype == 'table':
            if in_estrutura:
                continue
            add_md_table(doc, tok['lines'])
            consecutive_empty = 0

        # ---- Bloco de código ----------------------------------
        elif ttype == 'code':
            if in_estrutura:
                continue
            add_code_block(doc, tok['text'])
            consecutive_empty = 0

        # ---- Separador horizontal -----------------------------
        elif ttype == 'hr':
            consecutive_empty = 0
            continue  # ignora linhas ---

        # ---- Linha vazia --------------------------------------
        elif ttype == 'empty':
            consecutive_empty += 1
            if not in_estrutura and consecutive_empty <= 1:
                doc.add_paragraph(style=get_style(doc, 'BodyText'))

        # ---- Lista com marcador -------------------------------
        elif ttype == 'bullet':
            if in_estrutura:
                continue
            consecutive_empty = 0
            add_rich_paragraph(doc, tok['text'], style_id='ListParagraph')

        # ---- Lista numerada -----------------------------------
        elif ttype == 'numbered':
            if in_estrutura:
                continue
            consecutive_empty = 0
            add_rich_paragraph(doc, tok['text'], style_id='ListParagraph')

        # ---- Parágrafo normal ---------------------------------
        elif ttype == 'paragraph':
            if in_estrutura:
                continue
            consecutive_empty = 0
            add_rich_paragraph(doc, tok['text'], style_id='BodyText')

    return headings_h2


# ============================================================
# Fluxo principal para cada documento
# ============================================================

def create_document(task):
    md_path  = task['md']
    out_path = task['out']
    title    = task['title']

    print(f"\n  Processando: {md_path.name}")

    # 1. Copia o modelo (herda estilos, margens e layout)
    shutil.copy(str(MODELO), str(out_path))
    doc = Document(str(out_path))

    # 2. Limpa body preservando sectPr (margens)
    clear_body(doc)

    # 3. Lê e tokeniza o .md completo
    content = md_path.read_text(encoding='utf-8')
    tokens  = parse_md(content)

    # 4. Pré-coleta H2 headings para o índice (excluindo "Estrutura da Atividade")
    h2_headings = [
        t['text'] for t in tokens
        if t['type'] == 'heading2' and 'Estrutura da Atividade' not in t['text']
    ]

    # 5. Capa
    build_cover(doc, title)

    # 6. Quebra de página → índice
    add_page_break(doc)

    # 7. Índice
    build_index(doc, h2_headings)

    # 8. Conteúdo (renderiza tokens)
    render_tokens(doc, tokens)

    # 9. Salva
    doc.save(str(out_path))
    size = out_path.stat().st_size
    print(f"     -> Salvo: {out_path.name} ({size:,} bytes)")
    return out_path, size


# ============================================================
# Entry point
# ============================================================

if __name__ == '__main__':
    print("=" * 60)
    print("  CLAUDIO — Criando documentos ISG012")
    print("  TASK-ESTUDOS-AVANCADOS-2026-06-03 — Tarefa 2")
    print("=" * 60)

    results = []
    errors  = []

    for task in TASKS:
        try:
            path, size = create_document(task)
            results.append((path, size))
        except Exception as e:
            import traceback
            msg = f"ERRO em {task['md'].name}: {e}\n{traceback.format_exc()}"
            errors.append(msg)
            print(f"  [ERRO] {msg}")

    print("\n" + "=" * 60)
    print("  RESUMO FINAL")
    print("=" * 60)
    for path, size in results:
        print(f"  [OK] {path.name}")
        print(f"       {path}")
        print(f"       Tamanho: {size:,} bytes ({size / 1024:.1f} KB)")

    if errors:
        print(f"\n  ERROS ({len(errors)}):")
        for e in errors:
            print(f"  {e}")
    else:
        print(f"\n  Todos os {len(results)} documentos criados com sucesso.")

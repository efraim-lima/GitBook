"""
Mescla os Anexos preenchidos em um único arquivo .docx final.
Saída: Estágio Supervisionado - Efraim Lima.docx
"""

import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

BASE = Path(__file__).parent
PREENCHIDOS = BASE / "Preenchidos"
SAIDA = BASE / "Estágio Supervisionado - Efraim Lima.docx"

# Ordem dos anexos (a Declaração CLT é PDF — incluir como nota de página)
ANEXOS = [
    "Anexo_H_Capa.docx",
    "Anexo_D_PlanoAtividadesEstagio.docx",
    "Anexo_I_RelatorioParcial.docx",
    "Anexo_J_RelatorioFinal.docx",
    "Anexo_K_Autoavaliacao.docx",
    "Anexo_L_AvaliacaoSupervisor.docx",
    "Anexo_M_RequerimentoEquivalencia.docx",
    "Anexo_N_FichaEncerramento.docx",
]


def add_page_break(doc: Document):
    """Adiciona quebra de página no final do documento."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def main():
    # Verifica arquivos existentes
    faltando = [f for f in ANEXOS if not (PREENCHIDOS / f).exists()]
    if faltando:
        print(f"AVISO — arquivos não encontrados em Preenchidos/: {faltando}")

    presentes = [f for f in ANEXOS if (PREENCHIDOS / f).exists()]
    if not presentes:
        print("Nenhum arquivo encontrado em Preenchidos/. Abortando.")
        return

    print(f"Mesclando {len(presentes)} documentos...")

    # Documento base = primeiro anexo
    base_path = PREENCHIDOS / presentes[0]
    master = Document(str(base_path))
    print(f"  [base] {presentes[0]}")

    composer = Composer(master)

    for fname in presentes[1:]:
        path = PREENCHIDOS / fname
        print(f"  [+] {fname}")
        doc = Document(str(path))
        # Adiciona quebra de página antes de cada novo documento
        composer.append(doc)

    # Nota sobre a Declaração CLT (PDF separado)
    nota = master.add_paragraph()
    nota.add_run(
        "\n[NOTA: A Declaração de Experiência Profissional (CLT) é um documento "
        "PDF externo e deve ser inserida manualmente após o Anexo H (Capa) na "
        "versão impressa.]"
    )

    composer.save(str(SAIDA))
    size_kb = SAIDA.stat().st_size // 1024
    print(f"\nArquivo gerado: {SAIDA.name} ({size_kb} KB)")


if __name__ == "__main__":
    main()
